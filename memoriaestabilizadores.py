import win32com.client
from PIL import ImageGrab
import datetime
import tkinter as tk
from mailmerge import MailMerge
import locale
import re
import os
from tkinter import *
from tkinter import messagebox, ttk
from ttkthemes import ThemedStyle
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
from tkinter import filedialog
import PyPDF2



locale.setlocale(locale.LC_ALL, '')

class DocumentEditor:
    def __init__(self, document_path):
        self.document_path = document_path
        self.document = Document(self.document_path)

    def remove_empty_paragraphs_between_range(self, start_index, end_index):
        in_range = False

        paragraphs_to_remove = []

        for i, paragraph in enumerate(self.document.paragraphs):
            if i == start_index:
                in_range = True
            elif i == end_index + 1:
                in_range = False
            
            if in_range and not paragraph.text.strip():
                paragraphs_to_remove.append(paragraph)
        
        for paragraph in paragraphs_to_remove:
            p = paragraph._element
            p.getparent().remove(p)

    # para meter pdfs después de hacer el documento
    def append_pdf(self, input_pdf_path, output_path):
        pdf_writer = PyPDF2.PdfWriter()
        
        template_pdf_reader = PyPDF2.PdfReader(output_path)
        for page_num in range(len(template_pdf_reader.pages)):
            pdf_writer.add_page(template_pdf_reader.pages[page_num])

        for path in input_pdf_path:
            if path.lower().endswith('.pdf'):
                new_pdf_reader = PyPDF2.PdfReader(path)
                for page_num in range(len(new_pdf_reader.pages)):
                    pdf_writer.add_page(new_pdf_reader.pages[page_num])
            
        with open(output_path, 'wb') as output_file: 
            pdf_writer.write(output_file)

    # insertar planos en pdf antes del Apendice 1
    def in_planos(self, insert_pdf_paths, output_path):
        main_pdf = PyPDF2.PdfReader(output_path)

        insert_page = None
        for i in range(len(main_pdf.pages)):
            page = main_pdf.pages[i]
            text = page.extract_text()
            if "3. CARGA VIENTO" in text:
                insert_page = i
                break

        pdf_writer = PyPDF2.PdfWriter()

        if insert_page is None:
            print("Insert page not found")
        for i in range(insert_page):
            pdf_writer.add_page(main_pdf.pages[i])

        insert_pdf_paths = list(insert_pdf_paths)
        insert_pdf_paths.sort(key=lambda x: os.path.basename(x))
        
        for path in insert_pdf_paths:
            if path.lower().endswith('.pdf'):
                insert_pdf = PyPDF2.PdfReader(path)
                for i in range(len(insert_pdf.pages)):
                    pdf_writer.add_page(insert_pdf.pages[i])

        for i in range(insert_page, len(main_pdf.pages)):
            pdf_writer.add_page(main_pdf.pages[i])

        with open(output_path, 'wb') as output_file:
            pdf_writer.write(output_file)
    
    # Localizamos el index number para meter los TDS
    def buscar_txt_añTDS(self, texto_apendice):
        for i, paragraph in enumerate(self.document.paragraphs):
            if texto_apendice in paragraph.text:
                return i+1
        return -1

    def convert_to_pdf(word_file_path, pdf_file_path):
        convert(word_file_path, pdf_file_path)

    # Localizamos esto para meter la imagen de carga de viento
    def buscar_cargaviento(self, texto_viento):
        for i, paragraph in enumerate(self.document.paragraphs):
            if texto_viento in paragraph.text:
                return i+1
            return -1
        
    def añadir_img_viento(self, texto_viento, imagen_viento):
        target_index = self.buscar_txt_añTDS(texto_viento)
        if target_index != -1:
            target_paragraph = self.document.paragraphs[target_index]
            run = target_paragraph.add_run()
            run.add_picture(imagen_viento, width=Inches(7), height=Inches(9.2))
            return True 
        return False

    ########################## SuperSlim ########################################################

    # a�ade la imagen de la SS
    def buscar_txt_SS(self, texto_SS):
        for i, paragraph in enumerate(self.document.paragraphs):
            if texto_SS in paragraph.text:
                return i
        return -1

    def añadir_im_SS(self, texto_SS, imagen_SS):
        target_index = self.buscar_txt_SS(texto_SS)
        if target_index != -1:
            target_paragraph = self.document.paragraphs[target_index]
            run = target_paragraph.add_run()
            run.add_picture(imagen_SS, width=Inches(6.9), height=Inches(2.1))
            return True
        return False

    # añadir TDS del SS 
    def añadir_TDS_SS(self, texto_apendice, imagen_TDS_SS1, imagen_TDS_SS2, imagen_TDS_SS3, imagen_TDS_SS4, imagen_TDS_SS5, imagen_TDS_SS6, imagen_TDS_SS7, imagen_TDS_SS8, imagen_TDS_SS9, imagen_TDS_SS10, imagen_TDS_SS11, imagen_TDS_SS12, imagen_TDS_SS13, imagen_TDS_SS14, imagen_TDS_SS15, imagen_TDS_SS16, imagen_TDS_SS17, imagen_TDS_SS18, imagen_TDS_SS19, imagen_TDS_SS20, imagen_TDS_SS21, imagen_TDS_SS22, imagen_TDS_SS23, imagen_TDS_SS24, imagen_TDS_SS25, imagen_TDS_SS26, imagen_TDS_SS27, imagen_TDS_SS28):
        target_index = self.buscar_txt_añTDS(texto_apendice)
        if target_index != -1:
            target_paragraph = self.document.paragraphs[target_index]
            run = target_paragraph.add_run()
            run.add_picture(imagen_TDS_SS1, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS2, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS3, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS4, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS5, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS6, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS7, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS8, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS9, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS10, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS11, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS12, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS13, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS14, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS15, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS16, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS17, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS18, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS19, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS20, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS21, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS22, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS23, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS24, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS25, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS26, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS27, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS28, width=Inches(6.8), height=Inches(9.1))
            return True 
        return False


    #################################### Megaprop ################################################################

    # A�ade la imagen del MP
    def buscar_txt_MP(self, texto_MP):
        for i, paragraph in enumerate(self.document.paragraphs):
            if texto_MP in paragraph.text:
                return i
        return -1

    def añadir_im_MP(self, texto_MP, imagen_MP):
        target_index = self.buscar_txt_MP(texto_MP)
        if target_index != -1:
            target_paragraph = self.document.paragraphs[target_index]
            run = target_paragraph.add_run()
            run.add_picture(imagen_MP, width=Inches(6.9), height=Inches(2.1))
            return True
        return False

    # añadir TDS del Megaprop
    def añadir_TDS_MP(self, texto_apendice, imagen_TDS_MP1, imagen_TDS_MP2, imagen_TDS_MP3, imagen_TDS_MP4, imagen_TDS_MP5, imagen_TDS_MP6, imagen_TDS_MP7, imagen_TDS_MP8, imagen_TDS_MP9, imagen_TDS_MP10, imagen_TDS_MP11):
        target_index = self.buscar_txt_añTDS(texto_apendice)
        if target_index != -1:
            target_paragraph = self.document.paragraphs[target_index]
            run = target_paragraph.add_run()
            run.add_picture(imagen_TDS_MP1, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_MP2, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_MP3, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_MP4, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_MP5, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_MP6, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_MP7, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_MP8, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_MP9, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_MP10, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_MP11, width=Inches(6.8), height=Inches(9.1))
            return True 
        return False

    ########################################################## Granshor ###############################################################

    # A�ade la imagen del GS
    def buscar_txt_GS(self, texto_GS):
        for i, paragraph in enumerate(self.document.paragraphs):
            if texto_GS in paragraph.text:
                return i
        return -1

    def añadir_im_GS(self, texto_GS, imagen_GS):
        target_index = self.buscar_txt_GS(texto_GS)
        if target_index != -1:
            target_paragraph = self.document.paragraphs[target_index]
            run = target_paragraph.add_run()
            run.add_picture(imagen_GS, width=Inches(6.9), height=Inches(2.1))
            return True
        return False

    # añadir TDS del Granshor
    def añadir_TDS_GS(self, texto_apendice, imagen_TDS_GS1, imagen_TDS_GS2, imagen_TDS_GS3, imagen_TDS_GS4, imagen_TDS_GS5, imagen_TDS_GS6, imagen_TDS_GS7, imagen_TDS_GS8, imagen_TDS_GS9, imagen_TDS_GS10, imagen_TDS_GS11, imagen_TDS_GS12, imagen_TDS_GS13, imagen_TDS_GS14, imagen_TDS_GS15, imagen_TDS_GS16, imagen_TDS_GS17, imagen_TDS_GS18, imagen_TDS_GS19, imagen_TDS_GS20, imagen_TDS_GS21, imagen_TDS_GS22, imagen_TDS_GS23, imagen_TDS_GS24, imagen_TDS_GS25, imagen_TDS_GS26, imagen_TDS_GS27, imagen_TDS_GS28, imagen_TDS_GS29, imagen_TDS_GS30, imagen_TDS_GS31, imagen_TDS_GS32):
        target_index = self.buscar_txt_añTDS(texto_apendice)
        if target_index != -1:
            target_paragraph = self.document.paragraphs[target_index]
            run = target_paragraph.add_run()
            run.add_picture(imagen_TDS_GS1, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS2, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS3, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS4, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS5, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS6, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS7, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS8, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS9, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS10, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS11, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS12, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS13, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS14, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS15, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS16, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS17, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS18, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS19, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS20, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS21, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS22, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS23, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS24, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS25, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS26, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS27, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS28, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS29, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS30, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS31, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS32, width=Inches(6.8), height=Inches(4.5))
            return True 
        return False

    ##################################################### Lolashor ###################################################################

    # A�ade la imagen del Lolashor
    def buscar_txt_LS(self, texto_LS):
        for i, paragraph in enumerate(self.document.paragraphs):
            if texto_LS in paragraph.text:
                return i
        return -1

    def añadir_im_LS(self, texto_LS, imagen_LS):
        target_index = self.buscar_txt_LS(texto_LS)
        if target_index != -1:
            target_paragraph = self.document.paragraphs[target_index]
            run = target_paragraph.add_run()
            run.add_picture(imagen_LS, width=Inches(6.9), height=Inches(3.9))
            return True
        return False

    # añadir TDS del Tensor Cuadrado
    def añadir_TDS_LS(self, texto_apendice, imagen_TDS_LS1):
        target_index = self.buscar_txt_añTDS(texto_apendice)
        if target_index != -1:
            target_paragraph = self.document.paragraphs[target_index]
            run = target_paragraph.add_run()
            run.add_picture(imagen_TDS_LS1, width=Inches(6.8), height=Inches(2.5))
            return True 
        return False


    ################################################### Guardado del documento ######################################################################

    def save_document(self, output_path):
        self.document.save(output_path)
        pdf_output_path = os.path.splitext(output_path)[0] + ".pdf"
        convert(output_path, pdf_output_path)        


class Application(tk.Frame):
    def __init__(self, master=None):
        super().__init__(master)
        self.master = master
        self.master.title("Nota de Calculo de ESTABILIZADORES")
        self.master.geometry("880x1000")
        self.master.configure(background="#F5F5F5")
        self.pack(fill=tk.BOTH, expand=True)
        self.create_widgets()
        self.planos_paths = ()
        self.apendice_path = ()
                # Button styling 

    def create_widgets(self):

        # C�digo de Obra
        self.codigo_frame = tk.Frame(self, bg="#F5F5F5")
        self.codigo_frame.pack(pady=5)
        self.codigo_label = tk.Label(self.codigo_frame, text="Codigo de la obra:", font=("Helvetica", 14), bg="#F5F5F5", fg="#333333")
        self.codigo_label.pack(side=tk.LEFT, padx=5)
        self.codigo_entry = tk.Entry(self.codigo_frame, font=("Helvetica", 14))
        self.codigo_entry.pack(side=tk.RIGHT, padx=5, expand=True, fill=tk.X)

        # Obra
        self.obra_frame = tk.Frame(self, bg="#F5F5F5")
        self.obra_frame.pack(pady=5)
        self.obra_label = tk.Label(self.obra_frame, text="Obra:", font=("Helvetica", 14), bg="#F5F5F5", fg="#333333")
        self.obra_label.pack(side=tk.LEFT, padx=5)
        self.obra_entry = tk.Entry(self.obra_frame, font=("Helvetica", 14))
        self.obra_entry.pack(side=tk.RIGHT, padx=5, expand=True, fill=tk.X)

        # Direcci�n de la Obra
        self.dir_obra_frame = tk.Frame(self, bg="#F5F5F5")
        self.dir_obra_frame.pack(pady=5)
        self.Direccion_obra_label = tk.Label(self.dir_obra_frame, text="Direccion de la Obra:", font=("Helvetica", 14), bg="#F5F5F5", fg="#333333")
        self.Direccion_obra_label.pack(side=tk.LEFT, padx=5)
        self.Direccion_obra_entry = tk.Entry(self.dir_obra_frame, font=("Helvetica", 14))
        self.Direccion_obra_entry.pack(side=tk.RIGHT, padx=5, expand=True, fill=tk.X)

        # Cliente
        self.cliente_frame = tk.Frame(self, bg="#F5F5F5")
        self.cliente_frame.pack(pady=5)
        self.nombre_cliente_label = tk.Label(self.cliente_frame, text="Cliente:", font=("Helvetica", 14), bg="#F5F5F5", fg="#333333")
        self.nombre_cliente_label.pack(side=tk.LEFT, padx=5)
        self.nombre_cliente_entry = tk.Entry(self.cliente_frame, font=("Helvetica", 14))
        self.nombre_cliente_entry.pack(side=tk.RIGHT, padx=5, expand=True, fill=tk.X)

        # Altura de la fachada
        self.alt_frame = tk.Frame(self, bg="#F5F5F5")
        self.alt_frame.pack(pady=5)
        self.alt_label = tk.Label(self.alt_frame, text="Altura de la fachada:", font=("Helvetica", 14), bg="#F5F5F5", fg="#333333")
        self.alt_label.pack(side=tk.LEFT, padx=5)
        self.alt_entry = tk.Entry(self.alt_frame, font=("Helvetica", 14))
        self.alt_entry.pack(side=tk.RIGHT, padx=5, expand=True, fill=tk.X)

        # Ptje huecos en fachada 
        self.pctj_frame = tk.Frame(self, bg="#F5F5F5")
        self.pctj_frame.pack(pady=5)
        self.pctj_label = tk.Label(self.pctj_frame, text="Porcentaje de huecos en fachada:", font=("Helvetica", 14), bg="#F5F5F5", fg="#333333")
        self.pctj_label.pack(side=tk.LEFT, padx=5)
        self.pctj_entry = tk.Entry(self.pctj_frame, font=("Helvetica", 14))
        self.pctj_entry.pack(side=tk.RIGHT, padx=5, expand=True, fill=tk.X)

        # Coef Eólico 
        self.ceol_frame = tk.Frame(self, bg="#F5F5F5")
        self.ceol_frame.pack(pady=16)
        self.ceol_label = tk.Label(self.ceol_frame, text="Coeficiente eólico:", font=("Helvetica", 14), bg="#F5F5F5", fg="#333333")
        self.ceol_label.pack(side=tk.LEFT, padx=5)
        self.ceol_entry = tk.Entry(self.ceol_frame, font=("Helvetica", 14))
        self.ceol_entry.pack(side=tk.RIGHT, padx=5, expand=True, fill=tk.X)

        # Duracion de la obra
        self.dur_frame = tk.Frame(self, bg="#F5F5F5")
        self.dur_frame.pack(pady=7)
        self.dur_label = tk.Label(self.dur_frame, text="Duración de la obra:", font=("Helvetica", 14), bg="#F5F5F5", fg="#333333")
        self.dur_label.pack(side=tk.LEFT, padx=7)

        # Checkboxes
        self.checkboxes_frame3 = tk.Frame(self, bg="#F5F5F5")
        self.checkboxes_frame3.pack(pady=7)
        self.checkbar3 = Checkbar(self.checkboxes_frame3, ['Menor o igual a un año',  'Mayor a un año'], checkbox_font=("Helvetica", 14))
        self.checkbar3.pack(side=tk.TOP, fill=tk.X, padx=7)
        self.checkbar3.config(relief=tk.GROOVE, bd=4)

        # Seleccionar familia de materiales
        self.familia_frame = tk.Frame(self, bg="#F5F5F5")
        self.familia_frame.pack(pady=7)
        self.familia_label = tk.Label(self.familia_frame, text="Seleccionar tipo de obra:", font=("Helvetica", 14), bg="#F5F5F5", fg="#333333")
        self.familia_label.pack(side=tk.LEFT, padx=7)

        # Checkboxes
        self.checkboxes_frame = tk.Frame(self, bg="#F5F5F5")
        self.checkboxes_frame.pack(pady=7)
        self.checkbar = Checkbar(self.checkboxes_frame, ['Torre contrapesos',  'Torre encepados', 'Rigidizador muros medianeros', 'Rigidizador fachada'], checkbox_font=("Helvetica", 14))
        self.checkbar.pack(side=tk.TOP, fill=tk.X, padx=7)
        self.checkbar.config(relief=tk.GROOVE, bd=4)

        # Seleccionar familia de materiales
        self.familia_frame = tk.Frame(self, bg="#F5F5F5")
        self.familia_frame.pack(pady=7)
        self.familia_label = tk.Label(self.familia_frame, text="Seleccionar materiales utilizados:", font=("Helvetica", 14), bg="#F5F5F5", fg="#333333")
        self.familia_label.pack(side=tk.LEFT, padx=7)

        self.checkboxes_frame2 = tk.Frame(self, bg="#F5F5F5")
        self.checkboxes_frame2.pack(pady=7)
        self.checkbar2 = Checkbar(self.checkboxes_frame2, ['Superslim', 'Megaprop', 'Lolashor', 'Granshor'], checkbox_font=("Helvetica", 14))
        self.checkbar2.pack(side=tk.TOP, fill=tk.X, padx=7)
        self.checkbar2.config(relief=tk.GROOVE, bd=4)

        # Comboboxes
        self.combobox_frame = tk.Frame(self, bg="#F5F5F5")
        self.combobox_frame.pack(pady=7)

        self.label1 = ttk.Label(self.combobox_frame, text="Autor de la nota de cálculo", font=("Arial", 14))
        self.label1.grid(column=0, row=0, padx=11, pady=11)
        self.opcion_autor = tk.StringVar()
        opciones = ("José M. Maldonado", "David Lara.", "Ezequiel Sánchez.", "Andrés Rodríguez.", "Jorge Nebreda.", "Alberto Aldama.", "Adelaida Sáez.", "Alejandro Ángel Builes.", "Juan José Morón.", "Manuel González.", "Rafael Mansilla.")
        self.combobox_autor = ttk.Combobox(self.combobox_frame, width=20, textvariable=self.opcion_autor, values=opciones, font=("Arial", 12), style='Custom.TCombobox')
        self.combobox_autor.current(0)
        self.combobox_autor.grid(column=0, row=1, padx=11, pady=11)

        # Comboboxes button frame 
        self.label_revisor = ttk.Label(self.combobox_frame, text="Revisor de la nota de cálculo", font=("Arial", 14))
        self.label_revisor.grid(column=0, row=2, padx=11, pady=11)
        self.opcion_revisor = tk.StringVar()
        self.combobox_revisor = ttk.Combobox(self.combobox_frame, width=20, textvariable=self.opcion_revisor, values=opciones, font=("Arial", 12), style='Custom.TCombobox')
        self.combobox_revisor.current(0)
        self.combobox_revisor.grid(column=0, row=3, padx=11, pady=11)


        self.select_button = tk.Button(text="Ubicación de guardado:", command=self.select_output_path, font=("Helvetica", 16), bg="#F5F5F5", fg="black",
                               padx=17,
                               pady=9)
        self.select_button.pack()

        self.apendice_button = tk.Button(text="Seleccionar Excel Cálculos:", command=self.select_apendice, font=("Helvetica", 16), bg="#F5F5F5", fg="black",
                        padx=12,
                        pady=4)
        self.apendice_button.pack()

        self.apendice_button = tk.Button(text="Seleccionar PLANOS:", command=self.select_planos, font=("Helvetica", 16), bg="#F5F5F5", fg="black",
                        padx=12,
                        pady=4)
        self.apendice_button.pack()

        # Modificar button bg="#3986F3"
        self.fill_button = tk.Button(text="Crear", command=self.fill_template, font=("Helvetica", 16), bg="#FF6E40", fg="white",
                               padx=70,
                               pady=20)
        self.fill_button.pack()

    def select_apendice(self):
        folder = filedialog.askopenfilename()
        if folder:
            self.apendice_path = folder 

    def select_planos(self):
        folder = filedialog.askopenfilenames()
        if folder:
            self.planos_paths = folder
        
    def select_output_path(self):
        codigo = self.codigo_entry.get()
        current_date = datetime.datetime.now()
        folder = filedialog.askdirectory()
        if folder:
            folder = folder.replace('\\', '/')
            custom_date_format = current_date.strftime("%y%m%d")
            self.output_path = folder + f"/{codigo}_Nota_de_calculo_{custom_date_format}.docx"

    def fill_template(self):
        # Las entradas de texto
        nombre_cliente = self.nombre_cliente_entry.get()
        obra = self.obra_entry.get()
        Direccion_obra = self.Direccion_obra_entry.get()
        codigo = self.codigo_entry.get()
        selected_option2 = self.combobox_revisor.get()
        selected_option = self.combobox_autor.get()
        altura = self.alt_entry.get()
        huecos = self.pctj_entry.get()
        c_eolico = self.ceol_entry.get()

        wb_file_name = self.apendice_path
        outputPNGImage = 'C:/Memorias y servidor/Estabilizadores/carga_viento.jpg'
        xls_file = win32com.client.gencache.EnsureDispatch("Excel.Application")
        wb = xls_file.Workbooks.Open(Filename=wb_file_name)
        xls_file.DisplayAlerts = False 
        ws = wb.Worksheets("Estabilizador")
        ws.Range(ws.Cells(46,1),ws.Cells(110,14)).CopyPicture(Format= win32com.client.constants.xlBitmap)  # example from cell (1,1) to cell (15,3)
        img = ImageGrab.grabclipboard()
        img.save(outputPNGImage)
        wb.Close(SaveChanges=False, Filename=wb_file_name)
                
        additional_info = {
        "José M. Maldonado": "José Manuel Maldonado.\nMáster Ingeniero de Caminos, CC. y PP.\nDpto. Ingeniería INCYE.",
        "David Lara.": "David Lara.\nMáster Ingeniero de Caminos, CC. y PP.\nDpto. Ingeniería INCYE.",
        "Ezequiel Sánchez.": "Ezequiel Sánchez.\nIngeniero Industrial.\nDpto. Ingeniería INCYE.",
        "Andrés Rodríguez.": "Andrés Rodríguez Pérez.\nIngeniero Téc. Industrial\nDpto. Ingeniería INCYE.",
        "Jorge Nebreda.": "Jorge Nebreda.\nIngeniero de Caminos, CC. y PP.\nDpto. Ingeniería INCYE.",
        "Alberto Aldama.": "Alberto Aldama Martínez.\nIngeniero Industrial.\nDpto. Ingeniería INCYE.",
        "Adelaida Sáez.": "Adelaida Sáez Castejón.\nIng Téc. Industrial.\nDpto. Ingeniería INCYE.",
        "Alejandro Ángel Builes.": "Alejandro Ángel Builes.\nIngeniero Civil.\nDpto. Ingeniería INCYE.",
        "Juan José Morón.": "Juan José Morón Blanco.\nDelineante.\nDpto. Ingeniería INCYE.",
        "Manuel González.": "Manuel González-Arquiso Madrigal.\nIng. Téc. Agrícola.\nDpto. Ingeniería INCYE.",
        "Rafael Mansilla.": "Rafael Mansilla Correa.\nMáster Ingeniero de Caminos, CC. y PP.\nDpto. Ingeniería INCYE."
        }

        additional_info2 = {
        "José M. Maldonado": "JMM",
        "David Lara.": "DLM",
        "Ezequiel Sánchez.": "ESG",
        "Andrés Rodríguez.": "ARP",
        "Jorge Nebreda.": "JNS",
        "Alberto Aldama.": "AAM",
        "Adelaida Sáez.": "ASC",
        "Alejandro Ángel Builes.": "AAB",
        "Juan José Morón.": "JJM",
        "Manuel González.": "MGM",
        "Rafael Mansilla.": "RMC"
        }

        additional_info3 = {
        "José M. Maldonado": "Málaga",
        "David Lara.": "Madrid",
        "Ezequiel Sánchez.": "Madrid",
        "Andrés Rodríguez.": "Madrid",
        "Jorge Nebreda.": "Madrid",
        "Alberto Aldama.": "Bilbao",
        "Adelaida Sáez.": "Valencia",
        "Alejandro Ángel Builes.": "Madrid",
        "Juan José Morón.": "Sevilla",
        "Manuel González.": "Valladolid",
        "Rafael Mansilla.": "Madrid"
        }

        autor_nota = additional_info.get(selected_option, "")
        result = re.sub(' +', ' ', autor_nota)
        autor_nota = result

        revisor_nota = additional_info.get(selected_option2, "")
        result2 = re.sub(' +', ' ', revisor_nota)
        revisor_nota = result2   

        siglas_autor = additional_info2.get(selected_option, "")
        siglas_rev = additional_info2.get(selected_option2, "")
        ciudad = additional_info3.get(selected_option, "")

        # Fechas
        current_date = datetime.datetime.now()
        formatted_date = current_date.strftime("%d/%m/%Y")
        dia = current_date.strftime("%d")
        mes = current_date.strftime("%B")  # la B es el mes en formato palabra
        anyo = current_date.strftime("%Y") # mis primeras 

        # cargamos la plantilla
        template = "C:/Memorias y servidor/Estabilizadores/XXXXXXXXX_Nota de cálculo_211213 - copia2.docx"
        document = MailMerge(template)

        # Sustituimos valores
        document.merge(Nombre_Cliente=nombre_cliente, Obra=obra, Direccion_Obra=Direccion_obra, Codigo_Obra=codigo, Fecha=formatted_date, Dia=dia, Mes=mes, Anyo=anyo, Autor_NotaC=autor_nota, Revisor_NotaC=revisor_nota, Inic_AutorNC = siglas_autor, Inic_RevNC = siglas_rev, Ciudad=ciudad, Alt_Fach=altura, pctjehuecos=huecos, CoefEo=c_eolico)

        # Obtener los valores de las checkboxes
        checkbox_values = list(self.checkbar.state())
        checkbox_values2 = list(self.checkbar2.state())
        checkbox_values3 = list(self.checkbar3.state())

        # Guardar los valores en el documento
        document.merge(Pipeshor6=checkbox_values[0], Pipeshor4L=checkbox_values[1], Pipeshor4S=checkbox_values[2], Megaprop=checkbox_values[3])
        document.merge(INCYE300=checkbox_values2[0], INCYE450=checkbox_values2[1], INCYE600=checkbox_values2[2], SuperSlim=checkbox_values2[3])
        
        if checkbox_values[0]:   # Torre sobre contrapesos de hormigón
            document.merge(Torre_CH="La solución analizada es un estabilizador a base de torres ancladas a contrapesos de hormigón. \n\nMediante este sistema, las acciones horizontales (presión de viento y potencial desplome) se derivan de la fachada a las correas metálicas perimetrales, conectadas a una serie de torres metálicas. Dichas torres trabajan a cortante y flexión (convertidos en esfuerzos axiles al tratarse de celosías) y bajan las cargas hasta un contrapeso de hormigón en masa (no ejecutado por Incye), cuyas dimensiones y peso son tales que impiden el vuelco y deslizamiento del conjunto. \n\nDentro de las soluciones de refuerzo de fachadas, suelen ser los sistemas más robustos y con los que se consigue una baja solicitación en las fachadas al aportar las torres una rigidez similar a los pilares del edificio en su configuración de servicio. En esencia, las torres actúan como imanes de la carga (contrafuertes) y así alivian a las fachadas. \n\nPor otro lado, esta solución facilita los trabajos de demolición en el interior del edificio al minimizar la ocupación interior. \n\nPor el contrario, la presencia de bloques implica la ocupación de la vía adyacente al edificio, sumado al coste de encofrado, hormigonado, picado y vertido controlado del hormigón. \n\nFinalmente, la presencia de bloques de mucho peso concentrado en un espacio reducido requiere, por parte de un técnico competente, la comprobación de las tensiones transmitidas al terreno, así como los posibles asientos inducidos en las inmediaciones de la fachada, o los bulbos de presión horizontal comunicada a la base de la fachada. Estas cuestiones cobran especial relevancia si existen galerías subterráneas, o terrenos de baja capacidad portante, o nivel freático alto, o si se piensa realizar una excavación bajo rasante para la ejecución de sótanos o un recalce de la fachada.")
        if checkbox_values[1]:  # Torres sobre encepados
            document.merge(Encep="La solución analizada es un estabilizador a base de torres ancladas a encepados. \n\nMediante este sistema, las acciones horizontales (presión de viento y potencial desplome) se derivan de la fachada a las correas metálicas perimetrales, conectadas a una serie de torres metálicas. Dichas torres trabajan a cortante y flexión (convertidos en esfuerzos axiles al tratarse de celosías) y bajan las cargas hasta un encepado de pilotes o micropilotes (no ejecutado por Incye), los cuales mandan la carga a niveles más profundos y resistentes del terreno. \n\nDentro de las soluciones de refuerzo de fachadas, suelen estar entre los sistemas más robustos y con los que se consigue una baja solicitación en las fachadas al aportar las torres una rigidez similar a los pilares del edificio en su configuración de servicio. En esencia, las torres actúan como imanes de la carga (contrafuertes) y así alivian a las fachadas. \n\nPor otro lado, esta solución facilita los trabajos de demolición en el interior del edificio al minimizar la ocupación interior. \n\nPor el contrario, la presencia de encepados implica la ocupación de la vía adyacente al edificio, sumado al coste de ejecución y posterior demolición de los propios encepados. \n\nEl uso de torres sobre encepados es poco habitual, principalmente por su coste y complejidad. Los encepados pueden ser necesarios cuando el terreno subsuperficial no presenta una capacidad portante suficiente como para admitir la instalación de un bloque de hormigón en masa sobre la acera. Por ejemplo, en zonas con un nivel freático elevado, terrenos fangosos, o suelos con galerías subterráneas a poca profundidad.")
        if checkbox_values[2]:   # Rigidizador interior de fachada contra muros medianeros vecinos:
            document.merge(Rigid_murosvec="La solución analizada es un rigidizador o arriostramiento interior con puntales apoyados en medianeras. \n\nMediante este sistema, las acciones horizontales (presión de viento y potencial desplome) se derivan de la fachada a las correas metálicas perimetrales, y de estas a puntales horizontales que transmiten las cargas a los muros medianeros vecinos. \n\nDentro de las soluciones de refuerzo de fachadas, son sistemas robustos a la vez que livianos ya que los puntales permiten la transmisión horizontal de la carga desde la fachada a puntos fijos cercanos por el camino más corto posible. De este modo, cada nivel de rigidización simula un forjado discontinuo a base de puntales metálicos tal que las acciones horizontales no llegan a la base de la fachada, por lo que ésta no se ve apenas solicitada. \n\nComo aspecto adverso, es necesario trabajar más tiempo dentro del edificio para montar esta solución, lo cual puede condicionar el inicio del derribo de forjados. Asimismo, las tareas de demolición pueden derivar en daños en los puntales por impacto de cascotes. \n\nAsimismo, para poder transmitir carga a medianeras del propio edificio, o a las medianeras o forjados de la finca vecina, es necesario solicitar previamente los permisos oportunos y comprobar la aptitud estructural de las superficies de apoyo pues, al ser habitualmente de fábrica antigua, suelen tener muy baja resistencia a la ejecución de anclajes que va a trabajar a cortante y tracción. Dichos anclajes son necesarios para conectar los puntales a las medianeras. \n\nUna solución habitual que atenúa el potencial daño causado por los anclajes directos a muros medianeros es la disposición de una retícula de correas horizontales y velas verticales sobre dichos muros, de manera que los puntales de rigidización inciden contra las vigas metálicas, las cuales difunden la carga entre múltiples puntos.")
        if checkbox_values[3]:   # Rigidizador interior de fachada por vinculación de muros del propio edificio:
            document.merge(Rigid_murosedif="La solución analizada es un rigidizador o arriostramiento interior con puntales interconectando las fachadas del propio edificio. \n\nMediante este sistema, las acciones horizontales (presión de viento y potencial desplome) se derivan de una fachada a otras del mismo edificio a las que se halla conectada mediante puntales. De este modo, la acción sobre una fachada se reparte entre varias y los puntales simplemente actúan como bielas conectoras. \n\nEsta tipología es habitual en edificios cuyas fachadas se conservan y forman un polígono cerrado o, como mínimo, una U. De este modo, cada nivel de rigidización simula un forjado discontinuo (diafragma) a base de puntales metálicos. La presencia de fachadas que hacen esquina provoca que unas se comporten como contrafuertes de las otras, lo cual justifica la idea de comunicar las fachadas para derivar las cargas hacia muros en vuelta. La vinculación entre fachadas mediante puntales horizontales conforma células triangulares de alta rigidez que trabajan como un conjunto rígido. \n\nDentro de las soluciones de refuerzo de fachadas, esta tipología no transmite carga a un elemento más o menos fiable en cuanto a rigidez y capacidad portante, sino que hace trabajar a varias fachadas a la vez, a veces en su propio plano, a veces en perpendicular a su plano. \n\nPor tanto, es importante recalcar que, mediante esta solución, la carga se reparte entre varias fachadas y son éstas las responsables de transferir la carga al terreno. Por tanto, las fachadas asumen una solicitación de cortante y flexión probablemente superior a la que experimentarán con el edificio en su configuración de servicio (con forjados y pilares que bajan la carga horizontal a la cimentación), y superior también a la que se verían sometidas con otras soluciones de refuerzo de fachada, como pueden ser los estabilizadores con torres. Debido a esto, cobra especial importancia la caracterización de las propiedades mecánicas y resistentes de las fachadas rigidizadas, por parte de un técnico competente.")
        
        if checkbox_values2[0]:   # SuperSlim
            document.merge(SS="El sistema Superslim (área neta de la sección 19,64 cm2) es un sistema de perfilería constituido por vigas compuestas formadas por dos perfiles en C, unidos mediante presillas situadas en distintas posiciones que configuran un sistema de vigas modulares de gran versatilidad.")
        if checkbox_values2[1]:   # Megaprop
            document.merge(MP="El sistema Megaprop (área sección 58,45 cm2) es un sistema de perfilería constituido por vigas compuestas formadas por dos perfiles en C, unidos mediante presillas situadas en distintas posiciones que configuran un sistema de vigas modulares de gran versatilidad. El acero utilizado para su fabricación es de la calidad S355.")
        if checkbox_values2[2]:  # Granshor
            document.merge(GS="El sistema Granshor (área sección 72,59 cm2/cordón x 2 cordones) es un sistema de de celosías modular y sus elementos asociados. Fabricado con acero S355.")        
        if checkbox_values2[3]:  # Lolashor
            document.merge(Lola="texto del lolashor")


        ## textos de la metodología de cálculo 

        if checkbox_values2[0] and not checkbox_values2[1]:      # Texto del SS en la metodolog�a de cálculo
            document.merge(SST="Nota: En el Technical Data Sheet del sistema Superslim los valores incluidos en las gráficas son valores en Estado Límite de Servicio, es decir, son valores ya minorados por un coeficiente de 1,50. Para comparar frente a cargas mayoradas es necesario multiplicar los valores admisibles de las gráficas por 1,50 para no tener en cuenta un factor de mayoración duplicado.")
        if checkbox_values2[1] and not checkbox_values2[0]:      # Texto del MP en la metodolog�a de c�culo
            document.merge(SST="Nota: En el Technical Data Sheet del sistema Megaprop los valores incluidos en las gráficas son valores en Estado Límite de Servicio, es decir, son valores ya minorados por un coeficiente de 1,50. Para comparar frente a cargas mayoradas es necesario multiplicar los valores admisibles de las gráficas por 1,50 para no tener en cuenta un factor de mayoración duplicado.")
        if checkbox_values2[0] and checkbox_values2[1]:
            document.merge(SST="Nota: En el Technical Data Sheet del sistema Superslim, así como en el del Megaprop, los valores incluidos en las gráficas son valores en Estado Límite de Servicio, es decir, son valores ya minorados por un coeficiente de 1,50. Para comparar frente a cargas mayoradas es necesario multiplicar los valores admisibles de las gráficas por 1,50 para no tener en cuenta un factor de mayoración duplicado.")

        ## textos de la duracion y procedimiento de cálculo

        if checkbox_values3[0] and not checkbox_values3[1]:
            document.merge(duracion="menor")
            document.merge(duracion2="10")

        if not checkbox_values3[0] and checkbox_values3[1]:
            document.merge(duracion="mayor")
            document.merge(duracion2="50")
        if checkbox_values[1]:
            document.merge(textoproced="Con el momento de vuelco anterior se obtienen las dimensiones mínimas de los contrapesos para contrarrestar este vuelco.")

        # Ruta de guardado
        doc_modified = "C:/Memorias y servidor/Estabilizadores/Plantilla_estab.docx"
        document.write(doc_modified)

        if __name__ == '__main__':
            document_path = 'C:/Memorias y servidor/Estabilizadores/Plantilla_estab.docx'

            # texto para meter pdfs
            texto_apendice = "7. CARACTERÍSTICAS GEOMÉTRICAS Y MECÁNICAS PERFILES INCYE"
            
            # texto e imágenes del Superslim
            texto_SS = "El sistema Superslim (área neta de la sección 19,64 cm2) es un sistema de perfilería constituido por vigas compuestas formadas por dos perfiles en C, unidos mediante presillas situadas en distintas posiciones que configuran un sistema de vigas modulares de gran versatilidad."
            imagen_SS = "C:/Memorias y servidor/Aplicacion de Memorias/Imagenes/ss.JPG"
            
            # texto e imágenes del Megaprop
            texto_MP = "El sistema Megaprop (área sección 58,45 cm2) es un sistema de perfilería constituido por vigas compuestas formadas por dos perfiles en C, unidos mediante presillas situadas en distintas posiciones que configuran un sistema de vigas modulares de gran versatilidad. El acero utilizado para su fabricación es de la calidad S355."
            imagen_MP = "C:/Memorias y servidor/Aplicacion de Memorias/Imagenes/megaprop.JPG"
            
            # texto e imágenes del Granshor
            texto_GS = "El sistema Granshor (área sección 72,59 cm2/cordón x 2 cordones) es un sistema de de celosías modular y sus elementos asociados. Fabricado con acero S355."
            imagen_GS = "C:/Memorias y servidor/Aplicacion de Memorias/Imagenes/granshor.JPG"

            # Imagenes TDS del SuperSlim
            imagen_TDS_SS1 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-01.jpg"
            imagen_TDS_SS2 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-02.jpg"
            imagen_TDS_SS3 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-03.jpg"
            imagen_TDS_SS4 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-04.jpg"
            imagen_TDS_SS5 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-05.jpg"
            imagen_TDS_SS6 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-06.jpg"
            imagen_TDS_SS7 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-07.jpg"
            imagen_TDS_SS8 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-08.jpg"
            imagen_TDS_SS9 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-09.jpg"
            imagen_TDS_SS10 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-10.jpg"
            imagen_TDS_SS11 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-11.jpg"
            imagen_TDS_SS12 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-12.jpg"
            imagen_TDS_SS13 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-13.jpg"
            imagen_TDS_SS14 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-14.jpg"
            imagen_TDS_SS15 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-15.jpg"
            imagen_TDS_SS16 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-16.jpg"
            imagen_TDS_SS17 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-17.jpg"
            imagen_TDS_SS18 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-18.jpg"
            imagen_TDS_SS19 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-19.jpg"
            imagen_TDS_SS20 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-20.jpg"
            imagen_TDS_SS21 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-21.jpg"
            imagen_TDS_SS22 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-22.jpg"
            imagen_TDS_SS23 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-23.jpg"
            imagen_TDS_SS24 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-24.jpg"
            imagen_TDS_SS25 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-25.jpg"
            imagen_TDS_SS26 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-26.jpg"
            imagen_TDS_SS27 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-27.jpg"
            imagen_TDS_SS28 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-28.jpg"

            # añadir Imagenes TDS Granshor
            imagen_TDS_GS1 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Granshor/201116_Granshor_Anejos-01.jpg"
            imagen_TDS_GS2 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Granshor/201116_Granshor_Anejos-02.jpg"
            imagen_TDS_GS3 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Granshor/201116_Granshor_Anejos-03.jpg"
            imagen_TDS_GS4 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Granshor/201116_Granshor_Anejos-04.jpg"
            imagen_TDS_GS5 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Granshor/201116_Granshor_Anejos-05.jpg"
            imagen_TDS_GS6 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Granshor/201116_Granshor_Anejos-06.jpg"
            imagen_TDS_GS7 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Granshor/201116_Granshor_Anejos-07.jpg"
            imagen_TDS_GS8 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Granshor/201116_Granshor_Anejos-08.jpg"
            imagen_TDS_GS9 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Granshor/201116_Granshor_Anejos-09.jpg"
            imagen_TDS_GS10 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Granshor/201116_Granshor_Anejos-10.jpg"
            imagen_TDS_GS11 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Granshor/201116_Granshor_Anejos-11.jpg"
            imagen_TDS_GS12 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Granshor/201116_Granshor_Anejos-12.jpg"
            imagen_TDS_GS13 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Granshor/201116_Granshor_Anejos-13.jpg"
            imagen_TDS_GS14 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Granshor/201116_Granshor_Anejos-14.jpg"
            imagen_TDS_GS15 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Granshor/201116_Granshor_Anejos-15.jpg"
            imagen_TDS_GS16 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Granshor/201116_Granshor_Anejos-16.jpg"
            imagen_TDS_GS17 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Granshor/201116_Granshor_Anejos-17.jpg"
            imagen_TDS_GS18 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Granshor/201116_Granshor_Anejos-18.jpg"
            imagen_TDS_GS19 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Granshor/201116_Granshor_Anejos-19.jpg"
            imagen_TDS_GS20 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Granshor/201116_Granshor_Anejos-20.jpg"
            imagen_TDS_GS21 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Granshor/201116_Granshor_Anejos-21.jpg"
            imagen_TDS_GS22 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Granshor/201116_Granshor_Anejos-22.jpg"
            imagen_TDS_GS23 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Granshor/201116_Granshor_Anejos-23.jpg"
            imagen_TDS_GS24 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Granshor/201116_Granshor_Anejos-24.jpg"
            imagen_TDS_GS25 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Granshor/201116_Granshor_Anejos-25.jpg"
            imagen_TDS_GS26 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Granshor/201116_Granshor_Anejos-26.jpg"
            imagen_TDS_GS27 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Granshor/201116_Granshor_Anejos-27.jpg"
            imagen_TDS_GS28 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Granshor/201116_Granshor_Anejos-28.jpg"
            imagen_TDS_GS29 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Granshor/201116_Granshor_Anejos-29.jpg"
            imagen_TDS_GS30 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Granshor/201116_Granshor_Anejos-30.jpg"
            imagen_TDS_GS31 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Granshor/201116_Granshor_Anejos-31.jpg"
            imagen_TDS_GS32 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Granshor/201116_Granshor_Anejos-32.jpg"


            # Imagenes TDS Megaprop
            imagen_TDS_MP1 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Megaprop/TDS Megaprop INCYE-01.jpg"
            imagen_TDS_MP2 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Megaprop/TDS Megaprop INCYE-02.jpg"
            imagen_TDS_MP3 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Megaprop/TDS Megaprop INCYE-03.jpg"
            imagen_TDS_MP4 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Megaprop/TDS Megaprop INCYE-04.jpg"
            imagen_TDS_MP5 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Megaprop/TDS Megaprop INCYE-05.jpg"
            imagen_TDS_MP6 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Megaprop/TDS Megaprop INCYE-06.jpg"
            imagen_TDS_MP7 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Megaprop/TDS Megaprop INCYE-07.jpg"
            imagen_TDS_MP8 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Megaprop/TDS Megaprop INCYE-08.jpg"
            imagen_TDS_MP9 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Megaprop/TDS Megaprop INCYE-09.jpg"
            imagen_TDS_MP10 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Megaprop/TDS Megaprop INCYE-10.jpg"
            imagen_TDS_MP11 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Megaprop/TDS Megaprop INCYE-11.jpg"

            # Imagen del viento 
            imagen_viento = "C:/Memorias y servidor/Estabilizadores/carga_viento.jpg"
            texto_viento= "3. CARGA VIENTO"

            # Falta el TDS del Lolashor 

            document_editor = DocumentEditor(document_path)
            added_imagen_SS = document_editor.añadir_im_SS(texto_SS, imagen_SS)
            added_imagen_MP = document_editor.añadir_im_MP(texto_MP, imagen_MP)
            added_imagen_GS = document_editor.añadir_im_GS(texto_GS, imagen_GS)
            # falta el del lolashor

            if checkbox_values2[0]: # Superslim
                added_imagen_TDS_SS = document_editor.añadir_TDS_SS(texto_apendice, imagen_TDS_SS1, imagen_TDS_SS2, imagen_TDS_SS3, imagen_TDS_SS4, imagen_TDS_SS5, imagen_TDS_SS6, imagen_TDS_SS7, imagen_TDS_SS8, imagen_TDS_SS9, imagen_TDS_SS10, imagen_TDS_SS11, imagen_TDS_SS12, imagen_TDS_SS13, imagen_TDS_SS14, imagen_TDS_SS15, imagen_TDS_SS16, imagen_TDS_SS17, imagen_TDS_SS18, imagen_TDS_SS19, imagen_TDS_SS20, imagen_TDS_SS21, imagen_TDS_SS22, imagen_TDS_SS23, imagen_TDS_SS24, imagen_TDS_SS25, imagen_TDS_SS26, imagen_TDS_SS27, imagen_TDS_SS28)
            if checkbox_values2[1]: # Megaprop
                added_imagen_TDS_MP = document_editor.añadir_TDS_MP(texto_apendice, imagen_TDS_MP1, imagen_TDS_MP2, imagen_TDS_MP3, imagen_TDS_MP4, imagen_TDS_MP5, imagen_TDS_MP6, imagen_TDS_MP7, imagen_TDS_MP8, imagen_TDS_MP9, imagen_TDS_MP10, imagen_TDS_MP11)
            if checkbox_values2[3]: # Granshor
                added_imagen_TDS_GS = document_editor.añadir_TDS_GS(texto_apendice, imagen_TDS_GS1, imagen_TDS_GS2, imagen_TDS_GS3, imagen_TDS_GS4, imagen_TDS_GS5, imagen_TDS_GS6, imagen_TDS_GS7, imagen_TDS_GS8, imagen_TDS_GS9, imagen_TDS_GS10, imagen_TDS_GS11, imagen_TDS_GS12, imagen_TDS_GS13, imagen_TDS_GS14, imagen_TDS_GS15, imagen_TDS_GS16, imagen_TDS_GS17, imagen_TDS_GS18, imagen_TDS_GS19, imagen_TDS_GS20, imagen_TDS_GS21, imagen_TDS_GS22, imagen_TDS_GS23, imagen_TDS_GS24, imagen_TDS_GS25, imagen_TDS_GS26, imagen_TDS_GS27, imagen_TDS_GS28, imagen_TDS_GS29, imagen_TDS_GS30, imagen_TDS_GS31, imagen_TDS_GS32)
            # falta el del lolashor 

            if checkbox_values[0] or checkbox_values[1] or checkbox_values[2] or checkbox_values[3]:
                added_im_viento = document_editor.añadir_img_viento(texto_viento, imagen_viento)

            start_paragraph_index = 50
            end_paragraph_index = 65 

            if added_imagen_SS or added_imagen_MP or added_imagen_GS or added_imagen_TDS_SS or added_imagen_TDS_GS or added_imagen_TDS_MP or added_im_viento:
                if self.output_path:
                    document_editor.remove_empty_paragraphs_between_range(start_paragraph_index, end_paragraph_index)
                    
                    document_editor.save_document(self.output_path)
                    pdf_path = os.path.splitext(self.output_path)[0] + ".pdf"
                    document_editor.in_planos(self.planos_paths, pdf_path) 
                    #pdf_output_path = self.output_path.replace(".docx", ".pdf") 
                    #convert(self.output_path, pdf_output_path)
            else:
                print("The target paragraph was not found in the document. Image not added.")

        messagebox.showinfo("Completado!", "La nota de cálculo ha sido creada y guardada con éxito.")
        

class Checkbar(Frame):
    def __init__(self, parent=None, picks=[], side=LEFT, anchor=W, checkbox_font=None):
        Frame.__init__(self, parent)
        self.vars = []
        self.checkbox_font = checkbox_font if checkbox_font else ("Helvetica", 12)

        for pick in picks:
            var = IntVar()
            chk = Checkbutton(self, text=pick, variable=var, font=self.checkbox_font)
            chk.pack(side=side, anchor=anchor, expand=YES)
            self.vars.append(var)

    def state(self):
        return map((lambda var: var.get()), self.vars)


# Launch the UI
root = tk.Tk()
app = Application(master=root)
app.mainloop()




# este extracto de código hace una captura de pantalla de la hoja de excel en el rango A46-N110. Se puede usar para cualquier rango deseado
#wb_file_name = 'C:/Memorias y servidor/Estabilizadores/Granshor_210527.xlsx'
#outputPNGImage = 'C:/Memorias y servidor/Estabilizadores/test.jpg'

#xls_file = win32com.client.gencache.EnsureDispatch("Excel.Application")

#wb = xls_file.Workbooks.Open(Filename=wb_file_name)
#xls_file.DisplayAlerts = False 
#ws = wb.Worksheets("Estabilizador")
#ws.Range(ws.Cells(46,1),ws.Cells(110,14)).CopyPicture(Format= win32com.client.constants.xlBitmap)  # example from cell (1,1) to cell (15,3)
#img = ImageGrab.grabclipboard()
#img.save(outputPNGImage)
#wb.Close(SaveChanges=False, Filename=wb_file_name)