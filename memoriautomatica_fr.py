import datetime
import tkinter as tk
from mailmerge import MailMerge
import locale
import re
import os
from tkinter import *
from tkinter import messagebox, ttk
from ttkthemes import ThemedStyle
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
from tkinter import filedialog
import PyPDF2

locale.setlocale(locale.LC_ALL, 'fr_FR')

class DocumentEditorfr:
    def __init__(self, document_path):
        self.document_path = document_path
        self.document = Document(self.document_path)

    def remove_empty_paragraphs_between_range(self, start_index, end_index):
        in_range = False

        paragraphs_to_remove = []

        for i, paragraph in enumerate(self.document.paragraphs):
            if i == start_index:
                in_range = True
            elif i == end_index + 1:
                in_range = False
            
            if in_range and not paragraph.text.strip():
                paragraphs_to_remove.append(paragraph)
        
        for paragraph in paragraphs_to_remove:
            p = paragraph._element
            p.getparent().remove(p)

    def append_pdf(self, input_pdf_path, output_path):
        pdf_writer = PyPDF2.PdfWriter()
        
        template_pdf_reader = PyPDF2.PdfReader(output_path)
        for page_num in range(len(template_pdf_reader.pages)):
            pdf_writer.add_page(template_pdf_reader.pages[page_num])

        for path in input_pdf_path:
            if path.lower().endswith('.pdf'):
                new_pdf_reader = PyPDF2.PdfReader(path)
                for page_num in range(len(new_pdf_reader.pages)):
                    pdf_writer.add_page(new_pdf_reader.pages[page_num])
            
        with open(output_path, 'wb') as output_file: 
            pdf_writer.write(output_file)
    
    def in_planos(self, insert_pdf_paths, output_path):
        main_pdf = PyPDF2.PdfReader(output_path)

        insert_page = None
        for i in range(len(main_pdf.pages)):
            page = main_pdf.pages[i]
            text = page.extract_text()
            if "3. ANNEXES" in text:
                insert_page = i
                break

        pdf_writer = PyPDF2.PdfWriter()

        for i in range(insert_page):
            pdf_writer.add_page(main_pdf.pages[i])

        insert_pdf_paths = list(insert_pdf_paths)
        insert_pdf_paths.sort(key=lambda x: os.path.basename(x))
        
        for path in insert_pdf_paths:
            if path.lower().endswith('.pdf'):
                insert_pdf = PyPDF2.PdfReader(path)
                for i in range(len(insert_pdf.pages)):
                    pdf_writer.add_page(insert_pdf.pages[i])
                    i = i-1

        for i in range(insert_page, len(main_pdf.pages)):
            pdf_writer.add_page(main_pdf.pages[i])

        with open(output_path, 'wb') as output_file:
            pdf_writer.write(output_file)

    # Localizamos el index number para meter los TDS
    def buscar_txt_añTDS(self, texto_apendice):
        for i, paragraph in enumerate(self.document.paragraphs):
            if texto_apendice in paragraph.text:
                return i+1
        return -1
    
    def buscar_txt_temp(self, textotemp):
        for i, paragraph in enumerate(self.document.paragraphs):
            if textotemp in paragraph.text:
                return i+1
        return -1

    def convert_to_pdf(word_file_path, pdf_file_path):
        convert(word_file_path, pdf_file_path)

    ########################## SuperSlim ########################################################

    # a�ade la imagen de la SS
    def buscar_txt_SS(self, texto_SS):
        for i, paragraph in enumerate(self.document.paragraphs):
            if texto_SS in paragraph.text:
                return i
        return -1

    def añadir_im_SS(self, texto_SS, imagen_SS):
        target_index = self.buscar_txt_SS(texto_SS)
        if target_index != -1:
            target_paragraph = self.document.paragraphs[target_index]
            run = target_paragraph.add_run()
            run.add_picture(imagen_SS, width=Inches(6.9), height=Inches(2.1))
            return True
        return False

    # añadir TDS del SS 
    def añadir_TDS_SS(self, texto_apendice, imagen_TDS_SS1, imagen_TDS_SS2, imagen_TDS_SS3, imagen_TDS_SS4, imagen_TDS_SS5, imagen_TDS_SS6, imagen_TDS_SS7, imagen_TDS_SS8, imagen_TDS_SS9, imagen_TDS_SS10, imagen_TDS_SS11, imagen_TDS_SS12, imagen_TDS_SS13, imagen_TDS_SS14, imagen_TDS_SS15, imagen_TDS_SS16, imagen_TDS_SS17, imagen_TDS_SS18, imagen_TDS_SS19, imagen_TDS_SS20, imagen_TDS_SS21, imagen_TDS_SS22, imagen_TDS_SS23, imagen_TDS_SS24, imagen_TDS_SS25, imagen_TDS_SS26, imagen_TDS_SS27, imagen_TDS_SS28):
        target_index = self.buscar_txt_añTDS(texto_apendice)
        if target_index != -1:
            target_paragraph = self.document.paragraphs[target_index]
            run = target_paragraph.add_run()
            run.add_picture(imagen_TDS_SS1, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS2, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS3, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS4, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS5, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS6, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS7, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS8, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS9, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS10, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS11, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS12, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS13, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS14, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS15, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS16, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS17, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS18, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS19, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS20, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS21, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS22, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS23, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS24, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS25, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS26, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS27, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_SS28, width=Inches(6.8), height=Inches(9.1))
            return True 
        return False


    #################################### Megaprop ################################################################

    # A�ade la imagen del MP
    def buscar_txt_MP(self, texto_MP):
        for i, paragraph in enumerate(self.document.paragraphs):
            if texto_MP in paragraph.text:
                return i
        return -1

    def añadir_im_MP(self, texto_MP, imagen_MP):
        target_index = self.buscar_txt_MP(texto_MP)
        if target_index != -1:
            target_paragraph = self.document.paragraphs[target_index]
            run = target_paragraph.add_run()
            run.add_picture(imagen_MP, width=Inches(6.9), height=Inches(2.1))
            return True
        return False
    
    def añadir_im_tempMP(self, textotemp, imagentemp_MP):
        target_index = self.buscar_txt_temp(textotemp)
        if target_index != -1:
            target_paragraph = self.document.paragraphs[target_index]
            run = target_paragraph.add_run()
            run.add_picture(imagentemp_MP, width=Inches(4.9), height=Inches(0.55))
            return True
        return False

    # añadir TDS del Megaprop
    def añadir_TDS_MP(self, texto_apendice, imagen_TDS_MP1, imagen_TDS_MP2, imagen_TDS_MP3, imagen_TDS_MP4, imagen_TDS_MP5, imagen_TDS_MP6, imagen_TDS_MP7, imagen_TDS_MP8, imagen_TDS_MP9, imagen_TDS_MP10, imagen_TDS_MP11):
        target_index = self.buscar_txt_añTDS(texto_apendice)
        if target_index != -1:
            target_paragraph = self.document.paragraphs[target_index]
            run = target_paragraph.add_run()
            run.add_picture(imagen_TDS_MP1, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_MP2, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_MP3, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_MP4, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_MP5, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_MP6, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_MP7, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_MP8, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_MP9, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_MP10, width=Inches(6.8), height=Inches(9.1))
            run.add_picture(imagen_TDS_MP11, width=Inches(6.8), height=Inches(9.1))
            return True 
        return False
    
######################################## TDS de tensores telescópicos #####################################################################################################################

    def añadir_TDS_TT(self, texto_apendice, imagen_TDS_TT1):
        target_index = self.buscar_txt_añTDS(texto_apendice)
        if target_index != -1:
            target_paragraph = self.document.paragraphs[target_index]
            run = target_paragraph.add_run()
            run.add_picture(imagen_TDS_TT1, width=Inches(6.8), height=Inches(9.1))
            return True 
        return False


    ########################################################## Granshor ###############################################################

    # A�ade la imagen del GS
    def buscar_txt_GS(self, texto_GS):
        for i, paragraph in enumerate(self.document.paragraphs):
            if texto_GS in paragraph.text:
                return i
        return -1

    def añadir_im_GS(self, texto_GS, imagen_GS):
        target_index = self.buscar_txt_GS(texto_GS)
        if target_index != -1:
            target_paragraph = self.document.paragraphs[target_index]
            run = target_paragraph.add_run()
            run.add_picture(imagen_GS, width=Inches(6.9), height=Inches(2.1))
            return True
        return False
    
    def añadir_im_tempGS(self, textotemp, imagentemp_GS):
        target_index = self.buscar_txt_temp(textotemp)
        if target_index != -1:
            target_paragraph = self.document.paragraphs[target_index]
            run = target_paragraph.add_run()
            run.add_picture(imagentemp_GS, width=Inches(4.9), height=Inches(0.55))
            return True
        return False

    # añadir TDS del Granshor
    def añadir_TDS_GS(self, texto_apendice, imagen_TDS_GS1, imagen_TDS_GS2, imagen_TDS_GS3, imagen_TDS_GS4, imagen_TDS_GS5, imagen_TDS_GS6, imagen_TDS_GS7, imagen_TDS_GS8, imagen_TDS_GS9, imagen_TDS_GS10, imagen_TDS_GS11, imagen_TDS_GS12, imagen_TDS_GS13, imagen_TDS_GS14, imagen_TDS_GS15, imagen_TDS_GS16, imagen_TDS_GS17, imagen_TDS_GS18, imagen_TDS_GS19, imagen_TDS_GS20, imagen_TDS_GS21, imagen_TDS_GS22, imagen_TDS_GS23, imagen_TDS_GS24, imagen_TDS_GS25, imagen_TDS_GS26, imagen_TDS_GS27, imagen_TDS_GS28, imagen_TDS_GS29, imagen_TDS_GS30, imagen_TDS_GS31, imagen_TDS_GS32,imagen_TDS_GS33, imagen_TDS_GS34, imagen_TDS_GS35, imagen_TDS_GS36, imagen_TDS_GS37):
        target_index = self.buscar_txt_añTDS(texto_apendice)
        if target_index != -1:
            target_paragraph = self.document.paragraphs[target_index]
            run = target_paragraph.add_run()
            run.add_picture(imagen_TDS_GS1, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS2, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS3, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS4, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS5, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS6, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS7, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS8, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS9, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS10, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS11, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS12, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS13, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS14, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS15, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS16, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS17, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS18, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS19, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS20, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS21, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS22, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS23, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS24, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS25, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS26, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS27, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS28, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS29, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS30, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS31, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS32, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS33, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS34, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS35, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS36, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_GS37, width=Inches(6.8), height=Inches(4.5))
            return True 
        return False

    ##################################################### Tensor Cuadrado ###################################################################

    # A�ade la imagen del TC
    def buscar_txt_TC(self, texto_TC):
        for i, paragraph in enumerate(self.document.paragraphs):
            if texto_TC in paragraph.text:
                return i
        return -1

    def añadir_im_TC(self, texto_TC, imagen_TC):
        target_index = self.buscar_txt_GS(texto_TC)
        if target_index != -1:
            target_paragraph = self.document.paragraphs[target_index]
            run = target_paragraph.add_run()
            run.add_picture(imagen_TC, width=Inches(6.9), height=Inches(2.1))
            return True
        return False

    # añadir TDS del Tensor Cuadrado
    def añadir_TDS_TC(self, texto_apendice, imagen_TDS_TC1):
        target_index = self.buscar_txt_añTDS(texto_apendice)
        if target_index != -1:
            target_paragraph = self.document.paragraphs[target_index]
            run = target_paragraph.add_run()
            run.add_picture(imagen_TDS_TC1, width=Inches(6.8), height=Inches(4.5))
            return True 
        return False


    ############################################################ Pipeshor ########################################################

    # A�ade la imagen del PS4
    def buscar_txt_PS4(self, texto_PS4):
        for i, paragraph in enumerate(self.document.paragraphs):
            if texto_PS4 in paragraph.text:
                return i
        return -1

    def añadir_im_PS4(self, texto_PS4, imagen_PS4):
        target_index = self.buscar_txt_PS4(texto_PS4)
        if target_index != -1:
            target_paragraph = self.document.paragraphs[target_index]
            run = target_paragraph.add_run()
            run.add_picture(imagen_PS4, width=Inches(6.9), height=Inches(2.1))
            return True
        return False
    
    def buscar_txt_PS2(self, texto_PS2):
        for i, paragraph in enumerate(self.document.paragraphs):
            if texto_PS2 in paragraph.text:
                return i
        return -1

    def añadir_im_PS2(self, texto_PS2, imagen_PS2):
        target_index = self.buscar_txt_PS4(texto_PS2)
        if target_index != -1:
            target_paragraph = self.document.paragraphs[target_index]
            run = target_paragraph.add_run()
            run.add_picture(imagen_PS2, width=Inches(6.9), height=Inches(2.5))
            return True
        return False

    # A�ade imagen del PS6
    def buscar_txt_PS6(self, texto_PS6):
        for i, paragraph in enumerate(self.document.paragraphs):
            if texto_PS6 in paragraph.text:
                return i
        return -1

    def añadir_im_PS6(self, texto_PS6, imagen_PS6):
        target_index = self.buscar_txt_PS6(texto_PS6)
        if target_index != -1:
            target_paragraph = self.document.paragraphs[target_index]
            run = target_paragraph.add_run()
            run.add_picture(imagen_PS6, width=Inches(6.9), height=Inches(2.1))
            return True
        return False
    
    def añadir_im_tempP4L(self, textotemp, imagentemp_P4L):
        target_index = self.buscar_txt_temp(textotemp)
        if target_index != -1:
            target_paragraph = self.document.paragraphs[target_index]
            run = target_paragraph.add_run()
            run.add_picture(imagentemp_P4L, width=Inches(4.9), height=Inches(0.55))
            return True
        return False

    def añadir_im_tempP4S(self, textotemp, imagentemp_P4S):
        target_index = self.buscar_txt_temp(textotemp)
        if target_index != -1:
            target_paragraph = self.document.paragraphs[target_index]
            run = target_paragraph.add_run()
            run.add_picture(imagentemp_P4S, width=Inches(4.9), height=Inches(0.55))
            return True
        return False
    
    def añadir_im_tempP6(self, textotemp, imagentemp_P6):
        target_index = self.buscar_txt_temp(textotemp)
        if target_index != -1:
            target_paragraph = self.document.paragraphs[target_index]
            run = target_paragraph.add_run()
            run.add_picture(imagentemp_P6, width=Inches(4.9), height=Inches(0.55))
            return True
        return False

    # añadir TDS del Pipeshor
    def añadir_TDS_P(self, texto_apendice, imagen_TDS_P1, imagen_TDS_P2, imagen_TDS_P3, imagen_TDS_P4, imagen_TDS_P5, imagen_TDS_P6, imagen_TDS_P7, imagen_TDS_P8, imagen_TDS_P9, imagen_TDS_P10, imagen_TDS_P11, imagen_TDS_P12, imagen_TDS_P13, imagen_TDS_P14, imagen_TDS_P15, imagen_TDS_P16, imagen_TDS_P17, imagen_TDS_P18, imagen_TDS_P19, imagen_TDS_P20, imagen_TDS_P21, imagen_TDS_P22, imagen_TDS_P23, imagen_TDS_P24, imagen_TDS_P25, imagen_TDS_P26, imagen_TDS_P27, imagen_TDS_P28, imagen_TDS_P29, imagen_TDS_P30, imagen_TDS_P31, imagen_TDS_P32, imagen_TDS_P33, imagen_TDS_P34):
        target_index = self.buscar_txt_añTDS(texto_apendice)
        if target_index != -1:
            target_paragraph = self.document.paragraphs[target_index]
            run = target_paragraph.add_run()
            run.add_picture(imagen_TDS_P1, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_P2, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_P3, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_P4, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_P5, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_P6, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_P7, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_P8, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_P9, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_P10, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_P11, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_P12, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_P13, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_P14, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_P15, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_P16, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_P17, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_P18, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_P19, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_P20, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_P21, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_P22, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_P23, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_P24, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_P25, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_P26, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_P27, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_P28, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_P29, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_P30, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_P31, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_P32, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_P33, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_P34, width=Inches(6.8), height=Inches(4.5))
            return True 
        return False


    ######################################################## Vigas Incye #############################################################################

    # A�ade la imagen de las Vigas de reparto
    def buscar_txt_V(self, texto_V, texto_V2=None, texto_V3=None):
        for i, paragraph in enumerate(self.document.paragraphs):
            if texto_V in paragraph.text or (texto_V2 and texto_V2 in paragraph.text) or (texto_V3 and texto_V3 in paragraph.text):
                return i
        return -1

    def añadir_im_V(self, texto_V, imagen_V, texto_V2=None, texto_V3=None):
        target_index = self.buscar_txt_V(texto_V, texto_V2, texto_V3)
        if target_index != -1:
            target_paragraph = self.document.paragraphs[target_index]
            run = target_paragraph.add_run()
            run.add_picture(imagen_V, width=Inches(6.9), height=Inches(2.1))
            return True
        return False
    
    def buscar_txt_V2(self, texto_V2):
        for i, paragraph in enumerate(self.document.paragraphs):
            if texto_V2 in paragraph.text:
                return i
        return -1
    
    def añadir_im_V2(self, texto_V2, imagen_V2):
        target_index = self.buscar_txt_V2(texto_V2)
        if target_index != -1:
            target_paragraph = self.document.paragraphs[target_index]
            run = target_paragraph.add_run()
            run.add_picture(imagen_V2, width=Inches(6.9), height=Inches(2.1))
            return True
        return False

    # añadir TDS del INCYE 300
    def añadir_TDS_I3(self, texto_apendice, imagen_TDS_I31, imagen_TDS_I32):
        target_index = self.buscar_txt_añTDS(texto_apendice)
        if target_index != -1:
            target_paragraph = self.document.paragraphs[target_index]
            run = target_paragraph.add_run()
            run.add_picture(imagen_TDS_I31, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_I32, width=Inches(6.8), height=Inches(4.5))
            return True 
        return False

    # añadir TDS del INCYE 300
    def añadir_TDS_I4(self, texto_apendice, imagen_TDS_I41, imagen_TDS_I42):
        target_index = self.buscar_txt_añTDS(texto_apendice)
        if target_index != -1:
            target_paragraph = self.document.paragraphs[target_index]
            run = target_paragraph.add_run()
            run.add_picture(imagen_TDS_I41,width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_I42, width=Inches(6.8), height=Inches(4.5))
            return True 
        return False
    
    # añadir TDS del INCYE 300
    def añadir_TDS_I6(self, texto_apendice, imagen_TDS_I61, imagen_TDS_I62):
        target_index = self.buscar_txt_añTDS(texto_apendice)
        if target_index != -1:
            target_paragraph = self.document.paragraphs[target_index]
            run = target_paragraph.add_run()
            run.add_picture(imagen_TDS_I61, width=Inches(6.8), height=Inches(4.5))
            run.add_picture(imagen_TDS_I62, width=Inches(6.8), height=Inches(4.5))
            return True 
        return False


    ################################################### Guardado del documento ######################################################################

    def save_document(self, output_path):
        self.document.save(output_path)
        pdf_output_path = os.path.splitext(output_path)[0] + ".pdf"
        convert(output_path, pdf_output_path)        

class Applicationfr(tk.Frame):
    def __init__(self, master=None):
        super().__init__(master)
        self.master = master
        self.master.title("Nota de Calculo ACODALAMIENTO (FRANCIA)")
        self.master.geometry("880x930")
        self.master.configure(background="#F5F5F5")
        self.pack(fill=tk.BOTH, expand=True)
        self.create_widgets()
        self.planos_paths = ()
        self.apendice_paths = ()
                # Button styling 


    def create_widgets(self):

        # C�digo de Obra
        self.codigo_frame = tk.Frame(self, bg="#F5F5F5")
        self.codigo_frame.pack(pady=15)
        self.codigo_label = tk.Label(self.codigo_frame, text="Codigo de la obra:", font=("Helvetica", 14), bg="#F5F5F5", fg="#333333")
        self.codigo_label.pack(side=tk.LEFT, padx=15)
        self.codigo_entry = tk.Entry(self.codigo_frame, font=("Helvetica", 14))
        self.codigo_entry.pack(side=tk.RIGHT, padx=15, expand=True, fill=tk.X)

        # Obra
        self.obra_frame = tk.Frame(self, bg="#F5F5F5")
        self.obra_frame.pack(pady=15)
        self.obra_label = tk.Label(self.obra_frame, text="Obra:", font=("Helvetica", 14), bg="#F5F5F5", fg="#333333")
        self.obra_label.pack(side=tk.LEFT, padx=15)
        self.obra_entry = tk.Entry(self.obra_frame, font=("Helvetica", 14))
        self.obra_entry.pack(side=tk.RIGHT, padx=15, expand=True, fill=tk.X)

        # Direcci�n de la Obra
        self.dir_obra_frame = tk.Frame(self, bg="#F5F5F5")
        self.dir_obra_frame.pack(pady=15)
        self.Direccion_obra_label = tk.Label(self.dir_obra_frame, text="Direccion de la Obra:", font=("Helvetica", 14), bg="#F5F5F5", fg="#333333")
        self.Direccion_obra_label.pack(side=tk.LEFT, padx=15)
        self.Direccion_obra_entry = tk.Entry(self.dir_obra_frame, font=("Helvetica", 14))
        self.Direccion_obra_entry.pack(side=tk.RIGHT, padx=15, expand=True, fill=tk.X)

        # Cliente
        self.cliente_frame = tk.Frame(self, bg="#F5F5F5")
        self.cliente_frame.pack(pady=15)
        self.nombre_cliente_label = tk.Label(self.cliente_frame, text="Cliente:", font=("Helvetica", 14), bg="#F5F5F5", fg="#333333")
        self.nombre_cliente_label.pack(side=tk.LEFT, padx=15)
        self.nombre_cliente_entry = tk.Entry(self.cliente_frame, font=("Helvetica", 14))
        self.nombre_cliente_entry.pack(side=tk.RIGHT, padx=15, expand=True, fill=tk.X)

        # porcentaje de la carga axial
        self.ptje_frame = tk.Frame(self, bg="#F5F5F5")
        self.ptje_frame.pack(pady=10)
        self.ptje_label = tk.Label(self.ptje_frame, text="Porcentaje de carga axial:", font=("Helvetica", 14), bg="#F5F5F5", fg="#333333")
        self.ptje_label.pack(side=tk.LEFT, padx=15)
        self.ptje_entry = tk.Entry(self.ptje_frame, font=("Helvetica", 14))
        self.ptje_entry.pack(side=tk.RIGHT, padx=15, expand=True, fill=tk.X)

        # Seleccionar familia de materiales
        self.familia_frame = tk.Frame(self, bg="#F5F5F5")
        self.familia_frame.pack(pady=15)
        self.familia_label = tk.Label(self.familia_frame, text="Seleccionar familia de materiales:", font=("Helvetica", 14), bg="#F5F5F5", fg="#333333")
        self.familia_label.pack(side=tk.LEFT, padx=15)

        # Checkboxes
        self.checkboxes_frame = tk.Frame(self, bg="#F5F5F5")
        self.checkboxes_frame.pack(pady=15)
        self.checkbar = Checkbar(self.checkboxes_frame, ['INCYE 80', 'SuperSlim', 'Megaprop', 'Pipeshor 4L', 'Pipeshor 4S', 'Pipeshor 6'], checkbox_font=("Helvetica", 14))
        self.checkbar.pack(side=tk.TOP, fill=tk.X, padx=15)
        self.checkbar.config(relief=tk.GROOVE, bd=4)

        self.checkboxes_frame2 = tk.Frame(self, bg="#F5F5F5")
        self.checkboxes_frame2.pack(pady=15)
        self.checkbar2 = Checkbar(self.checkboxes_frame2, ['Granshor', 'INCYE 300', 'INCYE 450 SA', 'INCYE 450 TA', 'INCYE 600', 'T.Telescopico'], checkbox_font=("Helvetica", 14))
        self.checkbar2.pack(side=tk.TOP, fill=tk.X, padx=15)
        self.checkbar2.config(relief=tk.GROOVE, bd=4)

        # Comboboxes
        self.combobox_frame = tk.Frame(self, bg="#F5F5F5")
        self.combobox_frame.pack(pady=15)

        self.label1 = ttk.Label(self.combobox_frame, text="Autor de la nota de cálculo", font=("Arial", 14))
        self.label1.grid(column=0, row=0, padx=11, pady=11)
        self.opcion_autor = tk.StringVar()
        opciones = ("José M. Maldonado", "David Lara.", "Ezequiel Sánchez.", "Andrés Rodríguez.", "Jorge Nebreda.", "Alberto Aldama.", "Adelaida Sáez.", "Alejandro Ángel Builes.", "Juan José Morón.", "Manuel González.", "Rafael Mansilla.")
        self.combobox_autor = ttk.Combobox(self.combobox_frame, width=30, textvariable=self.opcion_autor, values=opciones, font=("Arial", 12), style='Custom.TCombobox')
        self.combobox_autor.current(0)
        self.combobox_autor.grid(column=0, row=1, padx=11, pady=11)

        # Comboboxes button frame 
        self.label_revisor = ttk.Label(self.combobox_frame, text="Revisor de la nota de cálculo", font=("Arial", 14))
        self.label_revisor.grid(column=0, row=2, padx=11, pady=11)
        self.opcion_revisor = tk.StringVar()
        self.combobox_revisor = ttk.Combobox(self.combobox_frame, width=30, textvariable=self.opcion_revisor, values=opciones, font=("Arial", 12), style='Custom.TCombobox')
        self.combobox_revisor.current(0)
        self.combobox_revisor.grid(column=0, row=3, padx=11, pady=11)

        self.select_button = tk.Button(text="Ubicación de guardado:", command=self.select_output_path, font=("Helvetica", 16), bg="#F5F5F5", fg="black",
                               padx=17,
                               pady=9)
        self.select_button.pack()

        self.apendice_button = tk.Button(text="Seleccionar Apéndice 2:", command=self.select_apendice, font=("Helvetica", 16), bg="#F5F5F5", fg="black",
                        padx=12,
                        pady=4)
        self.apendice_button.pack()

        self.apendice_button = tk.Button(text="Seleccionar PLANOS:", command=self.select_planos, font=("Helvetica", 16), bg="#F5F5F5", fg="black",
                        padx=12,
                        pady=4)
        self.apendice_button.pack()

        # Modificar button 
        self.fill_button = tk.Button(text="Crear", command=self.fill_templatefr, font=("Helvetica", 16), bg="#3986F3", fg="white",
                               padx=70,
                               pady=20)
        self.fill_button.pack()

    def select_apendice(self):
        folder = filedialog.askopenfilenames()
        if folder:
            self.apendice_paths = folder 

    def select_planos(self):
        folder = filedialog.askopenfilenames()
        if folder:
            self.planos_paths = folder

    def select_output_path(self):
        codigo = self.codigo_entry.get()
        current_date = datetime.datetime.now()
        folder = filedialog.askdirectory()
        if folder:
            folder = folder.replace('\\', '/')
            custom_date_format = current_date.strftime("%y%m%d")
            self.output_path = folder + f"/{codigo}_Note_de_calcul_{custom_date_format}.docx"

    def fill_templatefr(self):
        # Las entradas de texto
        nombre_cliente = self.nombre_cliente_entry.get()
        ptje_axial = self.ptje_entry.get()
        obra = self.obra_entry.get()
        Direccion_obra = self.Direccion_obra_entry.get()
        codigo = self.codigo_entry.get()
        selected_option2 = self.combobox_revisor.get()
        selected_option = self.combobox_autor.get()
        
        additional_info = {
        "José M. Maldonado": "José Manuel Maldonado.\nIngénieur Civil.\nDpt. Ingénierie INCYE.",
        "David Lara.": "David Lara.\nIngénieur Civil.\nDpt. Ingénierie INCYE.",
        "Ezequiel Sánchez.": "Ezequiel Sánchez.\nIngénieur industriel.\nDpto. Ingénierie INCYE.",
        "Andrés Rodríguez.": "Andrés Rodríguez Pérez.\nIng Tech. Industriel.\nDpt. Ingénierie INCYE.",
        "Jorge Nebreda.": "Jorge Nebreda.\nIngénieur Civil.\nDpt. Ingénierie INCYE.",
        "Alberto Aldama.": "Alberto Aldama Martínez.\nIngénieur Industriel.\nDpt. Ingénierie INCYE.",
        "Adelaida Sáez.": "Adelaida Sáez Castejón.\nIng Tech. Industriel.\nDpt. Ingénierie INCYE.",
        "Alejandro Ángel Builes.": "Alejandro Ángel Builes.\nIngénieur Civil.\nDpt. Ingénierie INCYE.",
        "Juan José Morón.": "Juan José Morón Blanco.\nDessinateur.\nDpt. Ingénierie INCYE.",
        "Manuel González.": "Manuel González-Arquiso Madrigal.\nIng. Tech. Agricole.\nDpt. Ingénierie INCYE.",
        "Rafael Mansilla.": "Rafael Mansilla Correa.\nIngénieur CIvil.\nDpt. Ingénierie INCYE."
        }

        additional_info2 = {
        "José M. Maldonado": "JMM",
        "David Lara.": "DLM",
        "Ezequiel Sánchez.": "ESG",
        "Andrés Rodríguez.": "ARP",
        "Jorge Nebreda.": "JNS",
        "Alberto Aldama.": "AAM",
        "Adelaida Sáez.": "ASC",
        "Alejandro Ángel Builes.": "AAB",
        "Juan José Morón.": "JJM",
        "Manuel González.": "MGM",
        "Rafael Mansilla.": "RMC"
        }

        additional_info3 = {
        "José M. Maldonado": "Málaga",
        "David Lara.": "Madrid",
        "Ezequiel Sánchez.": "Madrid",
        "Andrés Rodríguez.": "Madrid",
        "Jorge Nebreda.": "Madrid",
        "Alberto Aldama.": "Bilbao",
        "Adelaida Sáez.": "Valencia",
        "Alejandro Ángel Builes.": "Madrid",
        "Juan José Morón.": "Sevilla",
        "Manuel González.": "Valladolid",
        "Rafael Mansilla.": "Madrid"
        }

        autor_nota = additional_info.get(selected_option, "")
        result = re.sub(' +', ' ', autor_nota)
        autor_nota = result

        revisor_nota = additional_info.get(selected_option2, "")
        result2 = re.sub(' +', ' ', revisor_nota)
        revisor_nota = result2   

        siglas_autor = additional_info2.get(selected_option, "")
        siglas_rev = additional_info2.get(selected_option2, "")
        ciudad = additional_info3.get(selected_option, "")

        # Fechas
        current_date = datetime.datetime.now()
        formatted_date = current_date.strftime("%d/%m/%Y")
        dia = current_date.strftime("%d")
        mes = current_date.strftime("%B")  # la B es el mes en formato palabra
        anyo = current_date.strftime("%Y") # mis primeras 

        # cargamos la plantilla
        template = "C:/Memorias y servidor/Memorias Generadas/MemoriaIncyePruebafr13.docx"
        document = MailMerge(template)

        # Sustituimos valores
        document.merge(Nombre_Cliente=nombre_cliente, Obra=obra, Direccion_Obra=Direccion_obra, Codigo_Obra=codigo, pt_axial=ptje_axial, Fecha=formatted_date, Dia=dia, Mes=mes, Anyo=anyo, Autor_NotaC=autor_nota, Revisor_NotaC=revisor_nota, Inic_AutorNC = siglas_autor, Inic_RevNC = siglas_rev, Ciudad=ciudad)

        # Obtener los valores de las checkboxes
        checkbox_values = list(self.checkbar.state())
        checkbox_values2 = list(self.checkbar2.state())

        print("Checkbox values: ", checkbox_values)
        print("Checkbox values 2: ", checkbox_values2)

        # Guardar los valores en el documento
        document.merge(Pipeshor6=checkbox_values[0], Pipeshor4L=checkbox_values[1], Pipeshor4S=checkbox_values[2], Megaprop=checkbox_values[3])
        document.merge(INCYE300=checkbox_values2[0], INCYE450=checkbox_values2[1], INCYE600=checkbox_values2[2], SuperSlim=checkbox_values2[3])
        
        if checkbox_values[0]:   # INCYE80
            document.merge(TC="Le système INCYE 80 est un système modulaire constitué de tubes carrés de 80mm de côté et d'une épaisseur de 4mm. Ce système modulaire est compatible avec le système Superslim et ses éléments associés. Acier S275.")
        if checkbox_values[1]:   # SuperSlim
            document.merge(SS="Le système Superslim (aire nette de la section 19,64 cm2) est un système de profilés constitué de poutres composées formées par deux profilés en C. Il s’agit d’un système de poutres modulaires très polyvalent. Acier S355.")
        if checkbox_values[2]:   # Megaprop
            document.merge(MP="Le système Megaprop (aire de la section 58,45 cm2) est un système de profilés constitué de poutres composées par deux profilés en C. Elles sont attachées au moyen des platines situées à différentes positions qui configurent un système de poutres modulaires très polyvalent. L'acier utilisé dans sa fabrication est de qualité S355.")
        if checkbox_values[3]:   # Pipeshor 4L
            document.merge(Pip4L="Le système Pipeshor 4L, avec une aire de section de 100.13 cm2, est un système d'étais formé par modules de tubes de 406 mm de diamètre et leurs éléments associés. Fabriqué en acier S355 de 8 millimètres d'épaisseur.")
        if checkbox_values[4]:   # Pipeshor 4S
            document.merge(Pip4S="Le système Pipeshor 4S, avec une aire de section de 196,24 cm2, est un système d'étais formés par modules de tubes de 406 mm de diamètre et leurs éléments associés. Fabriqué en acier S355 de 16 millimètres d'épaisseur.")
        if checkbox_values[5]:  # Pipeshor 6
            document.merge(Pip6="Le système Pipeshor 6 (aire de la section 234,4 cm2) est formé par tubes de 610 mm de diamètre et leurs éléments associés. Fabriqué en acier de qualité S355 et d'une épaisseur de 12,5 millimètres.")
        if checkbox_values2[0]:  # Granshor
            document.merge(GS="Le système Granshor (aire de la section 72,59 cm2/corde x 2 cordes) est un système de treillis modulaire et ses éléments associés. Fabriqué en acier S355.")
        if checkbox_values2[1]:  # INCYE300
            document.merge(I300="Les liernes périmétrales sont exécutées avec des profilés INCYE300, constituées par poutres HEB300 renforcées avec une triple âme. Acier S275.")
        if checkbox_values2[2] or checkbox_values2[3]:  # INCYE450 SA o TA
            document.merge(I450="Les liernes périmétrales seront exécutées avec des profilés INCYE450, renforcées. Il y en existe deux types de profilés dans cette catégorie, l’uns avec une triple âme et les autres de simple âme avec des raidisseurs au niveau des points d’impact des butons pour éviter le phénomène de voilement.  Acier S275.")
        if checkbox_values2[4]:  # INCYE600
            document.merge(I600="Les liernes périmétrales seront exécutées avec de poutres INCYE600, constituées de poutres HEB600 renforcées renforcés avec une triple âme. Fabriqués en acier S275.")
        if checkbox_values[1]:      # Texto del SS en la metodolog�a de cálculo
            document.merge(SST="Nota: En el Technical Data Sheet del sistema Superslim los valores incluidos en las gráficas son valores en Estado Límite de Servicio, es decir, son valores ya minorados por un coeficiente de 1,50. Para comparar frente a cargas mayoradas es necesario multiplicar los valores admisibles de las gráficas por 1,50 para no tener en cuenta un factor de mayoración duplicado.")
        if checkbox_values[2]:      # Texto del MP en la metodolog�a de c�culo
            document.merge(MPT="Nota: En el Technical Data Sheet del sistema Megaprop los valores incluidos en las gráficas son valores en Estado Límite de Servicio, es decir, son valores ya minorados por un coeficiente de 1,50. Para comparar frente a cargas mayoradas es necesario multiplicar los valores admisibles de las gráficas por 1,50 para no tener en cuenta un factor de mayoración duplicado.")
        if checkbox_values2[1] or checkbox_values2[2] or checkbox_values2[3] or checkbox_values2[4]:   # Texto de los perfiles en la metodolog�a de cálculo
            document.merge(PERF="Para la comprobación de los perfiles se obtendrán los esfuerzos en cada uno de ellos y se compararán con sus valores admisibles.")

        # Ruta de guardado
        doc_modified = "C:/Memorias y servidor/Memorias Generadas/Modificada_MemoriaEjemplofr.docx"
        document.write(doc_modified)

        if __name__ == '__main__':
            document_path = 'C:/Memorias y servidor/Memorias Generadas/Modificada_MemoriaEjemplofr.docx'
            
            # texto e imágenes del Superslim
            texto_SS = "Le système Superslim (aire nette de la section 19,64 cm2) est un système de profilés constitué de poutres composées formées par deux profilés en C. Il s’agit d’un système de poutres modulaires très polyvalent. Acier S355."
            imagen_SS = "C:/Memorias y servidor/Aplicacion de Memorias/Imagenes/ss.JPG"
            # "C:/Memorias y servidor/Aplicacion de Memorias/Imagenes/ss.JPG"
            
            # texto e imágenes del Megaprop
            texto_MP = "Le système Megaprop (aire de la section 58,45 cm2) est un système de profilés constitué de poutres composées par deux profilés en C. Elles sont attachées au moyen des platines situées à différentes positions qui configurent un système de poutres modulaires très polyvalent. L'acier utilisé dans sa fabrication est de qualité S355."
            imagen_MP = "C:/Memorias y servidor/Aplicacion de Memorias/Imagenes/megaprop.JPG"
            
            # texto e imágenes del Granshor
            texto_GS = "Le système Granshor (aire de la section 72,59 cm2/corde x 2 cordes) est un système de treillis modulaire et ses éléments associés. Fabriqué en acier S355."
            imagen_GS = "C:/Memorias y servidor/Aplicacion de Memorias/Imagenes/granshor.JPG"
            
            # textos e imágenes del TC
            texto_TC = "Le système INCYE 80 est un système modulaire constitué de tubes carrés de 80mm de côté et d'une épaisseur de 4mm. Ce système modulaire est compatible avec le système Superslim et ses éléments associés. Acier S275."
            imagen_TC = "C:/Memorias y servidor/Aplicacion de Memorias/Imagenes/tensorcuadrado.JPG"

            # textos e imágenes del Pipeshor
            texto_PS6 = "Le système Pipeshor 6 (aire de la section 234,4 cm2) est formé par tubes de 610 mm de diamètre et leurs éléments associés. Fabriqué en acier de qualité S355 et d'une épaisseur de 12,5 millimètres."
            imagen_PS6 = "C:/Memorias y servidor/Aplicacion de Memorias/Imagenes/pipeshor6.JPG"
            texto_PS2 = "Le système Pipeshor 4S, avec une aire de section de 196,24 cm2, est un système d'étais formés par modules de tubes de 406 mm de diamètre et leurs éléments associés. Fabriqué en acier S355 de 16 millimètres d'épaisseur."
            texto_PS4 = "Le système Pipeshor 4L, avec une aire de section de 100.13 cm2, est un système d'étais formé par modules de tubes de 406 mm de diamètre et leurs éléments associés. Fabriqué en acier S355 de 8 millimètres d'épaisseur."
            imagen_PS4 = "C:/Memorias y servidor/Aplicacion de Memorias/Imagenes/pipeshor.JPG"
            imagen_PS2 = "C:/Memorias y servidor/Aplicacion de Memorias/Imagenes/pipeshors.JPG"

            # texto e imágenes de los INCYE
            texto_V = "Les liernes périmétrales sont exécutées avec des profilés INCYE300, constituées par poutres HEB300 renforcées avec une triple âme. Acier S275."
            texto_V2 = "Les liernes périmétrales seront exécutées avec des profilés INCYE450, renforcées. Il y en existe deux types de profilés dans cette catégorie, l’uns avec une triple âme et les autres de simple âme avec des raidisseurs au niveau des points d’impact des butons pour éviter le phénomène de voilement.  Acier S275."
            texto_V3 = "Les liernes périmétrales seront exécutées avec de poutres INCYE600, constituées de poutres HEB600 renforcées renforcés avec une triple âme. Fabriqués en acier S275."
            imagen_V = "C:/Memorias y servidor/Aplicacion de Memorias/Imagenes/perfil.JPG"
            imagen_V2 = "C:/Memorias y servidor/Aplicacion de Memorias/Imagenes/perfilSA.JPG"
            texto_apendice = "3. ANNEXES"

            # texto para meter las fórmulas de la temperatura
            textotemp = "Si nous considérons une jonction rigide et parfaite entre toutes les pièces des butons et entre les butons et l’écran, l’effort axial additionnel du à la température est le suivant :"
            
            # Imagenes TDS del SuperSlim
            imagen_TDS_SS1 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-01.jpg"
            imagen_TDS_SS2 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-02.jpg"
            imagen_TDS_SS3 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-03.jpg"
            imagen_TDS_SS4 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-04.jpg"
            imagen_TDS_SS5 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-05.jpg"
            imagen_TDS_SS6 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-06.jpg"
            imagen_TDS_SS7 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-07.jpg"
            imagen_TDS_SS8 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-08.jpg"
            imagen_TDS_SS9 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-09.jpg"
            imagen_TDS_SS10 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-10.jpg"
            imagen_TDS_SS11 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-11.jpg"
            imagen_TDS_SS12 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-12.jpg"
            imagen_TDS_SS13 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-13.jpg"
            imagen_TDS_SS14 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-14.jpg"
            imagen_TDS_SS15 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-15.jpg"
            imagen_TDS_SS16 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-16.jpg"
            imagen_TDS_SS17 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-17.jpg"
            imagen_TDS_SS18 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-18.jpg"
            imagen_TDS_SS19 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-19.jpg"
            imagen_TDS_SS20 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-20.jpg"
            imagen_TDS_SS21 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-21.jpg"
            imagen_TDS_SS22 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-22.jpg"
            imagen_TDS_SS23 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-23.jpg"
            imagen_TDS_SS24 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-24.jpg"
            imagen_TDS_SS25 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-25.jpg"
            imagen_TDS_SS26 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-26.jpg"
            imagen_TDS_SS27 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-27.jpg"
            imagen_TDS_SS28 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/SS/ANEJO SS-28.jpg"
            
            # Imagenes del TDS del INCYE 300
            imagen_TDS_I31 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/INCYE300/210427_Joints liernes_page-0009.jpg"
            imagen_TDS_I32 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/INCYE300/210427_Joints liernes_page-0010.jpg"

            # Imagenes del TDS del INCYE 450
            imagen_TDS_I41 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/INCYE450/210427_Joints liernes_page-0014.jpg"
            imagen_TDS_I42 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/INCYE450/210427_Joints liernes_page-0015.jpg"

            # Imagenes del TDS del INCYE 600
            imagen_TDS_I61 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/INCYE600/210427_Joints liernes_page-0019.jpg"
            imagen_TDS_I62 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/INCYE600/210427_Joints liernes_page-0020.jpg"

            # Imagenes TDS del Tensor Cuadrado
            imagen_TDS_TC1 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/TC/Tendeur carre 80_page-0001.jpg"

            # Imagenes TDS Pipeshor
            imagen_TDS_P1 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/PS_fr/200903_Pipeshor_francais_T_rafa_page-0001.jpg"
            imagen_TDS_P2 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/PS_fr/200903_Pipeshor_francais_T_rafa_page-0002.jpg"
            imagen_TDS_P3= "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/PS_fr/200903_Pipeshor_francais_T_rafa_page-0003.jpg"
            imagen_TDS_P4= "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/PS_fr/200903_Pipeshor_francais_T_rafa_page-0004.jpg"
            imagen_TDS_P5= "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/PS_fr/200903_Pipeshor_francais_T_rafa_page-0005.jpg"
            imagen_TDS_P6= "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/PS_fr/200903_Pipeshor_francais_T_rafa_page-0006.jpg"
            imagen_TDS_P7= "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/PS_fr/200903_Pipeshor_francais_T_rafa_page-0007.jpg"
            imagen_TDS_P8= "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/PS_fr/200903_Pipeshor_francais_T_rafa_page-0008.jpg"
            imagen_TDS_P9= "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/PS_fr/200903_Pipeshor_francais_T_rafa_page-0009.jpg"
            imagen_TDS_P10= "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/PS_fr/200903_Pipeshor_francais_T_rafa_page-0010.jpg"
            imagen_TDS_P11= "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/PS_fr/200903_Pipeshor_francais_T_rafa_page-0011.jpg"
            imagen_TDS_P12= "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/PS_fr/200903_Pipeshor_francais_T_rafa_page-0012.jpg"
            imagen_TDS_P13= "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/PS_fr/200903_Pipeshor_francais_T_rafa_page-0013.jpg"
            imagen_TDS_P14= "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/PS_fr/200903_Pipeshor_francais_T_rafa_page-0014.jpg"
            imagen_TDS_P15= "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/PS_fr/200903_Pipeshor_francais_T_rafa_page-0015.jpg"
            imagen_TDS_P16= "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/PS_fr/200903_Pipeshor_francais_T_rafa_page-0016.jpg"
            imagen_TDS_P17= "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/PS_fr/200903_Pipeshor_francais_T_rafa_page-0017.jpg"
            imagen_TDS_P18= "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/PS_fr/200903_Pipeshor_francais_T_rafa_page-0018.jpg"
            imagen_TDS_P19= "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/PS_fr/200903_Pipeshor_francais_T_rafa_page-0019.jpg"
            imagen_TDS_P20= "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/PS_fr/200903_Pipeshor_francais_T_rafa_page-0020.jpg"
            imagen_TDS_P21= "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/PS_fr/200903_Pipeshor_francais_T_rafa_page-0021.jpg"
            imagen_TDS_P22= "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/PS_fr/200903_Pipeshor_francais_T_rafa_page-0022.jpg"
            imagen_TDS_P23= "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/PS_fr/200903_Pipeshor_francais_T_rafa_page-0023.jpg"
            imagen_TDS_P24= "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/PS_fr/200903_Pipeshor_francais_T_rafa_page-0024.jpg"
            imagen_TDS_P25= "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/PS_fr/200903_Pipeshor_francais_T_rafa_page-0025.jpg"
            imagen_TDS_P26= "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/PS_fr/200903_Pipeshor_francais_T_rafa_page-0026.jpg"
            imagen_TDS_P27= "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/PS_fr/200903_Pipeshor_francais_T_rafa_page-0027.jpg"
            imagen_TDS_P28= "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/PS_fr/200903_Pipeshor_francais_T_rafa_page-0028.jpg"
            imagen_TDS_P29= "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/PS_fr/200903_Pipeshor_francais_T_rafa_page-0029.jpg"
            imagen_TDS_P30= "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/PS_fr/200903_Pipeshor_francais_T_rafa_page-0030.jpg"
            imagen_TDS_P31= "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/PS_fr/200903_Pipeshor_francais_T_rafa_page-0031.jpg"
            imagen_TDS_P32= "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/PS_fr/200903_Pipeshor_francais_T_rafa_page-0032.jpg"
            imagen_TDS_P33= "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/PS_fr/200903_Pipeshor_francais_T_rafa_page-0033.jpg"
            imagen_TDS_P34= "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/PS_fr/200903_Pipeshor_francais_T_rafa_page-0034.jpg"


            # añadir Imagenes TDS Granshor   "C:\Memorias y servidor\Aplicacion de Memorias\TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0001.jpg"
            imagen_TDS_GS1 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0001.jpg"
            imagen_TDS_GS2 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0002.jpg"
            imagen_TDS_GS3 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0003.jpg"
            imagen_TDS_GS4 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0004.jpg"
            imagen_TDS_GS5 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0005.jpg"
            imagen_TDS_GS6 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0006.jpg"
            imagen_TDS_GS7 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0007.jpg"
            imagen_TDS_GS8 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0008.jpg"
            imagen_TDS_GS9 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0009.jpg"
            imagen_TDS_GS10 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0010.jpg"
            imagen_TDS_GS11 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0011.jpg"
            imagen_TDS_GS12 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0012.jpg"
            imagen_TDS_GS13 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0013.jpg"
            imagen_TDS_GS14 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0014.jpg"
            imagen_TDS_GS15 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0015.jpg"
            imagen_TDS_GS16 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0016.jpg"
            imagen_TDS_GS17 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0017.jpg"
            imagen_TDS_GS18 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0018.jpg"
            imagen_TDS_GS19 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0019.jpg"
            imagen_TDS_GS20 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0020.jpg"
            imagen_TDS_GS21 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0021.jpg"
            imagen_TDS_GS22 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0022.jpg"
            imagen_TDS_GS23 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0023.jpg"
            imagen_TDS_GS24 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0024.jpg"
            imagen_TDS_GS25 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0025.jpg"
            imagen_TDS_GS26 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0026.jpg"
            imagen_TDS_GS27 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0027.jpg"
            imagen_TDS_GS28 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0028.jpg"
            imagen_TDS_GS29 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0029.jpg"
            imagen_TDS_GS30 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0030.jpg"
            imagen_TDS_GS31 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0031.jpg"
            imagen_TDS_GS32 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0032.jpg"
            imagen_TDS_GS33 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0033.jpg"
            imagen_TDS_GS34 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0034.jpg"
            imagen_TDS_GS35 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0035.jpg"
            imagen_TDS_GS36 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0036.jpg"
            imagen_TDS_GS37 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/GS_fr/Granshor 2.0.2_fr_page-0037.jpg"


            # Imagenes TDS Megaprop
            imagen_TDS_MP1 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Megaprop/TDS Megaprop INCYE-01.jpg"
            imagen_TDS_MP2 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Megaprop/TDS Megaprop INCYE-02.jpg"
            imagen_TDS_MP3 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Megaprop/TDS Megaprop INCYE-03.jpg"
            imagen_TDS_MP4 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Megaprop/TDS Megaprop INCYE-04.jpg"
            imagen_TDS_MP5 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Megaprop/TDS Megaprop INCYE-05.jpg"
            imagen_TDS_MP6 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Megaprop/TDS Megaprop INCYE-06.jpg"
            imagen_TDS_MP7 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Megaprop/TDS Megaprop INCYE-07.jpg"
            imagen_TDS_MP8 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Megaprop/TDS Megaprop INCYE-08.jpg"
            imagen_TDS_MP9 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Megaprop/TDS Megaprop INCYE-09.jpg"
            imagen_TDS_MP10 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Megaprop/TDS Megaprop INCYE-10.jpg"
            imagen_TDS_MP11 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs/Megaprop/TDS Megaprop INCYE-11.jpg"

            # Imágenes de las fórmulas
            imagentemp_MP = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/tempMP.jpg"
            imagentemp_GS = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/tempGS.jpg"
            imagentemp_P4L = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/tempP4L.jpg"
            imagentemp_P4S = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/tempP4S.jpg"
            imagentemp_P6 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/tempP6.jpg"

            # Imagen TDS Telescopico
            imagen_TDS_TT1 = "C:/Memorias y servidor/Aplicacion de Memorias/TDSs_fr/Telescopico/Tendeur telescopique_page-0001.jpg"


            document_editor = DocumentEditorfr(document_path)
            added_imagen_SS = document_editor.añadir_im_SS(texto_SS, imagen_SS)
            added_imagen_MP = document_editor.añadir_im_MP(texto_MP, imagen_MP)
            added_imagen_GS = document_editor.añadir_im_GS(texto_GS, imagen_GS)
            added_imagen_TC = document_editor.añadir_im_TC(texto_TC, imagen_TC)
            added_imagen_PS4 = document_editor.añadir_im_PS4(texto_PS4, imagen_PS4)
            added_imagen_PS2 = document_editor.añadir_im_PS2(texto_PS2, imagen_PS2)
            added_imagen_V = document_editor.añadir_im_V(texto_V, imagen_V, texto_V2, texto_V3)
            added_imagen_PS6 = document_editor.añadir_im_PS6(texto_PS6, imagen_PS6)

            paragraph_index = document_editor.buscar_txt_temp(textotemp)
            print("Paragraph index:", paragraph_index)
            
            if checkbox_values[2]: # Megaprop fórmula
                added_image_tempMP = document_editor.añadir_im_tempMP(textotemp, imagentemp_MP)
                print("image added", added_image_tempMP)            
            if checkbox_values[3]: # Pipeshor 4L fórmula
                added_image_tempP4L = document_editor.añadir_im_tempP4L(textotemp, imagentemp_P4L)
                print("image added", added_image_tempP4L)
            if checkbox_values[4]: # Pipeshor 4S fórmula
                added_image_tempP4S = document_editor.añadir_im_tempP4S(textotemp, imagentemp_P4S)
                print("image added", added_image_tempP4S)
            if checkbox_values[5]: # Pipeshor 6 fórmula
                added_image_tempP6 = document_editor.añadir_im_tempP6(textotemp, imagentemp_P6)
                print("image added", added_image_tempP6)
            if checkbox_values2[0]: # Granshor fórmula
                added_image_tempGS = document_editor.añadir_im_tempGS(textotemp, imagentemp_GS)
                print("image added", added_image_tempGS)            
            if checkbox_values[0]: # Tensor cuadrado / INCYE 80
                added_imagen_TDS_TC = document_editor.añadir_TDS_TC(texto_apendice, imagen_TDS_TC1)
            if checkbox_values[1]: # Superslim
                added_imagen_TDS_SS = document_editor.añadir_TDS_SS(texto_apendice, imagen_TDS_SS1, imagen_TDS_SS2, imagen_TDS_SS3, imagen_TDS_SS4, imagen_TDS_SS5, imagen_TDS_SS6, imagen_TDS_SS7, imagen_TDS_SS8, imagen_TDS_SS9, imagen_TDS_SS10, imagen_TDS_SS11, imagen_TDS_SS12, imagen_TDS_SS13, imagen_TDS_SS14, imagen_TDS_SS15, imagen_TDS_SS16, imagen_TDS_SS17, imagen_TDS_SS18, imagen_TDS_SS19, imagen_TDS_SS20, imagen_TDS_SS21, imagen_TDS_SS22, imagen_TDS_SS23, imagen_TDS_SS24, imagen_TDS_SS25, imagen_TDS_SS26, imagen_TDS_SS27, imagen_TDS_SS28)
            if checkbox_values[2]: # Megaprop
                added_imagen_TDS_MP = document_editor.añadir_TDS_MP(texto_apendice, imagen_TDS_MP1, imagen_TDS_MP2, imagen_TDS_MP3, imagen_TDS_MP4, imagen_TDS_MP5, imagen_TDS_MP6, imagen_TDS_MP7, imagen_TDS_MP8, imagen_TDS_MP9, imagen_TDS_MP10, imagen_TDS_MP11)    
            if checkbox_values2[0]: # Granshor
                added_imagen_TDS_GS = document_editor.añadir_TDS_GS(texto_apendice, imagen_TDS_GS1, imagen_TDS_GS2, imagen_TDS_GS3, imagen_TDS_GS4, imagen_TDS_GS5, imagen_TDS_GS6, imagen_TDS_GS7, imagen_TDS_GS8, imagen_TDS_GS9, imagen_TDS_GS10, imagen_TDS_GS11, imagen_TDS_GS12, imagen_TDS_GS13, imagen_TDS_GS14, imagen_TDS_GS15, imagen_TDS_GS16, imagen_TDS_GS17, imagen_TDS_GS18, imagen_TDS_GS19, imagen_TDS_GS20, imagen_TDS_GS21, imagen_TDS_GS22, imagen_TDS_GS23, imagen_TDS_GS24, imagen_TDS_GS25, imagen_TDS_GS26, imagen_TDS_GS27, imagen_TDS_GS28, imagen_TDS_GS29, imagen_TDS_GS30, imagen_TDS_GS31, imagen_TDS_GS32, imagen_TDS_GS33, imagen_TDS_GS34, imagen_TDS_GS35, imagen_TDS_GS36, imagen_TDS_GS37)
            if checkbox_values2[1]: # INCYE 300
                added_imagen_TDS_I3 = document_editor.añadir_TDS_I3(texto_apendice, imagen_TDS_I31, imagen_TDS_I32)
            if checkbox_values2[2]: # INCYE 450
                added_imagen_TDS_I4 = document_editor.añadir_TDS_I4(texto_apendice, imagen_TDS_I41, imagen_TDS_I42)
            if checkbox_values2[3]: # INCYE 300
                added_imagen_TDS_I6 = document_editor.añadir_TDS_I6(texto_apendice, imagen_TDS_I61, imagen_TDS_I62)
            if checkbox_values[3] or checkbox_values[4] or checkbox_values[5]: # Pipeshor
                added_image_TDS_P = document_editor.añadir_TDS_P(texto_apendice, imagen_TDS_P1, imagen_TDS_P2, imagen_TDS_P3, imagen_TDS_P4, imagen_TDS_P5, imagen_TDS_P6, imagen_TDS_P7, imagen_TDS_P8, imagen_TDS_P9, imagen_TDS_P10, imagen_TDS_P11, imagen_TDS_P12, imagen_TDS_P13, imagen_TDS_P14, imagen_TDS_P15, imagen_TDS_P16, imagen_TDS_P17, imagen_TDS_P18, imagen_TDS_P19, imagen_TDS_P20, imagen_TDS_P21, imagen_TDS_P22, imagen_TDS_P23, imagen_TDS_P24, imagen_TDS_P25, imagen_TDS_P26, imagen_TDS_P27, imagen_TDS_P28, imagen_TDS_P29, imagen_TDS_P30, imagen_TDS_P31, imagen_TDS_P32, imagen_TDS_P33, imagen_TDS_P34)           
            if checkbox_values2[5]:
                added_image_TDS_TT = document_editor.añadir_TDS_TT(texto_apendice, imagen_TDS_TT1)


            # asignamos el valor False a todas al comienzo para que no tomen valor directamente
            added_image_tempMP = False
            added_image_tempP4L = False
            added_image_tempP4S = False
            added_image_tempP6 = False
            added_image_tempGS = False

            start_paragraph_index = 39
            end_paragraph_index = 50

            if (added_image_tempMP or added_image_tempP4L or added_image_tempP4S or added_image_tempP6 or added_image_tempGS or added_imagen_SS or added_imagen_MP or added_imagen_GS or added_imagen_TC or added_imagen_PS4 or added_imagen_PS2 or added_imagen_V or added_imagen_PS6 or added_imagen_TDS_SS or added_imagen_TDS_I3  or added_imagen_TDS_I4 or added_imagen_TDS_I6 or added_imagen_TDS_TC or added_image_TDS_P or added_imagen_TDS_GS or added_imagen_TDS_MP or added_image_TDS_TT):
                if self.output_path:
                    document_editor.remove_empty_paragraphs_between_range(start_paragraph_index, end_paragraph_index)
                    document_editor.save_document(self.output_path)
                    pdf_path = os.path.splitext(self.output_path)[0] + ".pdf"
                    document_editor.in_planos(self.planos_paths, pdf_path) 
                    document_editor.append_pdf(self.apendice_paths, pdf_path) 
            else:
                print("The target paragraph was not found in the document. Image not added.")

        messagebox.showinfo("Completado", "El documento se ha modificado y guardado con éxito.")
        

class Checkbar(Frame):
    def __init__(self, parent=None, picks=[], side=LEFT, anchor=W, checkbox_font=None):
        Frame.__init__(self, parent)
        self.vars = []
        self.checkbox_font = checkbox_font if checkbox_font else ("Helvetica", 12)

        for pick in picks:
            var = IntVar()
            chk = Checkbutton(self, text=pick, variable=var, font=self.checkbox_font)
            chk.pack(side=side, anchor=anchor, expand=YES)
            self.vars.append(var)

    def state(self):
        return map((lambda var: var.get()), self.vars)


# Launch the UI
root = tk.Tk()
app = Applicationfr(master=root)
app.mainloop()


